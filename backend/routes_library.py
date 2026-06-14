"""
Library blueprint for Sh'elah.

Library tree, text-pane, search, word-meaning, export, and texts-index routes
extracted verbatim from ``app.py`` (Stage 2 blueprint split). Logic is
unchanged; only the route decorator target moved from ``@app.route`` to
``@routes_library.route`` and shared helpers/constants are imported from
``app``.
"""

import io
import re

from flask import Blueprint, jsonify, request, send_file

from backend.helpers import (
    QUICK_TEXT_ALIASES,
    COMMUNITIES,
    COMMUNITY_ALIASES,
    _decode_route_ref,
    _coerce_int,
    _extract_search_metadata_filters,
    _fill_missing_english_lines,
    _contains_hebrew_letters,
    _lookup_hebrew_word_meaning,
    _lookup_english_word_meaning,
    _translate_english_text_online,
    _collect_word_meaning_alternatives,
)
from backend.auth import maybe_require_clerk_auth

routes_library = Blueprint("library", __name__)


@routes_library.route("/api/library/index")
def library_index():
    """Returns report-adjusted Sefaria library tree (non-loading removals pruned, fix refs applied)."""
    from backend.sefaria_library import get_library_index
    data = get_library_index()
    return jsonify(data)


@routes_library.route("/api/library/leaf-refs")
def library_leaf_refs():
    """Return leaf refs for a given index title to power section-grid selectors."""
    import re as _re
    from backend.sefaria_library import get_index_entry, get_index_leaf_refs

    def _parse_talmud_daf_key(daf_str):
        """Convert a daf string like '2a', '13b' to a sortable integer key."""
        m = _re.search(r'(\d+)([ab])', str(daf_str or '').strip().lower())
        if not m:
            return None
        num = int(m.group(1))
        side = 0 if m.group(2) == 'a' else 1
        return num * 2 + side

    def _extract_talmud_sections(index_title, entry):
        """
        Extract named chapter sections from a Talmud index entry's alts.Chapters.
        Returns list of {label, fromDaf, toDaf} dicts, or [] if unavailable.
        """
        alts = entry.get("alts") if isinstance(entry, dict) else None
        if not isinstance(alts, dict):
            return []
        chapters_alt = alts.get("Chapters") or alts.get("chapters")
        if not isinstance(chapters_alt, dict):
            return []
        nodes = chapters_alt.get("nodes")
        if not isinstance(nodes, list):
            return []

        sections = []
        for node in nodes:
            if not isinstance(node, dict):
                continue
            raw_title = str(node.get("title") or "").strip()
            canonical_title = entry.get("title", index_title)
            whole_ref = str(node.get("wholeRef") or "").strip()

            # Parse the daf range from wholeRef: "Berakhot 2a:1-13a:15"
            # Strip the title prefix then read the daf range
            ref_body = whole_ref
            if canonical_title and ref_body.lower().startswith(canonical_title.lower()):
                ref_body = ref_body[len(canonical_title):].lstrip(" ,")
            elif index_title and ref_body.lower().startswith(index_title.lower()):
                ref_body = ref_body[len(index_title):].lstrip(" ,")

            range_m = _re.search(
                r'(\d+[ab])(?:[\d:]*)\s*-\s*(\d+[ab])', ref_body, _re.IGNORECASE)
            if not range_m:
                continue
            from_daf = range_m.group(1).lower()
            to_daf = range_m.group(2).lower()

            # Clean the chapter label: "Chapter 1; MeEimatai" → "MeEimatai (2a–13a)"
            # Keep the human name after the semicolon, if present
            if ';' in raw_title:
                label = raw_title.split(';', 1)[1].strip()
            else:
                label = raw_title

            sections.append({
                "label": label,
                "fromDaf": from_daf,
                "toDaf": to_daf,
            })

        return sections

    def _collapse_talmud_leaf_refs(index_title, refs, max_items=260):
        """Convert segment-level Talmud refs into unique daf refs for stable grid rendering."""
        if not isinstance(refs, list) or not refs:
            return []

        normalized_title = str(index_title or "").strip()
        compact_refs = []
        seen = set()

        for ref_value in refs:
            ref_text = str(ref_value or "").strip()
            if not ref_text:
                continue

            body = ref_text
            if normalized_title and body.lower().startswith(normalized_title.lower()):
                body = body[len(normalized_title):].lstrip(" ,")

            daf_match = _re.search(r"(\d+[ab])", body, _re.IGNORECASE)
            if not daf_match:
                continue

            daf = daf_match.group(1).lower()
            if daf in seen:
                continue
            seen.add(daf)

            compact_refs.append(f"{normalized_title} {daf}".strip())
            if len(compact_refs) >= max_items:
                break

        return compact_refs

    def _synthesize_section_refs(index_title, max_items=140):
        entry = get_index_entry(index_title)
        schema = entry.get("schema", {}) if isinstance(entry, dict) else {}
        if not isinstance(schema, dict):
            return [], []

        lengths = schema.get("lengths") if isinstance(
            schema.get("lengths"), list) else []
        if not lengths:
            return [], []

        try:
            first_level_count = int(lengths[0])
        except (TypeError, ValueError):
            return [], []

        if first_level_count <= 1:
            return [], []

        section_names = schema.get("sectionNames") if isinstance(
            schema.get("sectionNames"), list) else []
        address_types = schema.get("addressTypes") if isinstance(
            schema.get("addressTypes"), list) else []

        first_section_name = str(section_names[0] or "").strip(
        ).lower() if section_names else ""
        first_address_type = str(address_types[0] or "").strip(
        ).lower() if address_types else ""

        refs = []
        if first_section_name == "daf" or first_address_type == "talmud":
            # Sefaria Talmud indexing starts at 2a.
            for idx in range(first_level_count):
                daf_num = (idx // 2) + 2
                side = "a" if idx % 2 == 0 else "b"
                refs.append(f"{index_title} {daf_num}{side}")
                if len(refs) >= max_items:
                    break
            sections = _extract_talmud_sections(index_title, entry)
            return refs, sections

        for idx in range(1, first_level_count + 1):
            refs.append(f"{index_title} {idx}")
            if len(refs) >= max_items:
                break

        return refs, []

    requested_title = _decode_route_ref(request.args.get("title", ""))
    title = str(requested_title or "").strip()
    max_refs = _coerce_int(request.args.get("max"), 140,
                           min_value=1, max_value=260)

    if not title:
        return jsonify({"title": "", "refs": [], "sections": []})

    sections = []
    try:
        refs = get_index_leaf_refs(title, max_refs=max_refs)
    except Exception:
        refs = []

    if not isinstance(refs, list):
        refs = []

    if len(refs) <= 1:
        refs, sections = _synthesize_section_refs(title, max_items=max_refs)
    else:
        # Try to extract sections even when refs came from get_index_leaf_refs
        try:
            entry = get_index_entry(title)
            sections = _extract_talmud_sections(title, entry)
        except Exception:
            sections = []

    if sections:
        collapsed_refs = _collapse_talmud_leaf_refs(
            title, refs, max_items=max_refs)
        if collapsed_refs:
            refs = collapsed_refs

    return jsonify({
        "title": title,
        "refs": refs,
        "sections": sections,
    })


@routes_library.route("/api/library/popular")
def library_popular():
    """Returns curated popular texts per category."""
    from backend.sefaria_library import get_popular_texts
    return jsonify(get_popular_texts())


@routes_library.route("/api/text/<path:ref>")
def get_text_inline(ref):
    """Fetches a Sefaria text inline — Hebrew + English + metadata."""
    from backend.sefaria_library import get_text
    decoded_ref = _decode_route_ref(ref)
    data = get_text(decoded_ref)

    if isinstance(data, dict) and data.get("error"):
        error_type = data.get("error_type", "")
        if error_type == "sefaria_blocked" or "blocked" in str(data.get("error", "")).lower():
            return jsonify(data), 503

    should_translate = str(request.args.get("autotranslate", "1")).strip().lower() not in {
        "0", "false", "no", "off"
    }
    if should_translate and isinstance(data, dict) and not data.get("error"):
        data = _fill_missing_english_lines(data)

    return jsonify(data)


@routes_library.route("/api/diagnostics/sefaria")
def sefaria_diagnostics():
    """Real-time availability probe for the upstream Sefaria API.
    Returns status for both the v3 and v2 endpoints, plus cached block info.
    """
    from backend.sefaria_library import check_sefaria_availability
    result = check_sefaria_availability()
    http_status = 200 if result.get("overall_available") else 503
    return jsonify(result), http_status


@routes_library.route("/api/word/meaning")
def get_word_meaning():
    """Look up a highlighted word meaning (best-effort for Hebrew and English)."""
    raw_word = str(request.args.get("word", "") or "").strip()
    if not raw_word:
        return jsonify({"error": "Missing word parameter"}), 400

    requested_lang = str(request.args.get(
        "lang", "en") or "en").strip().lower()
    if requested_lang not in {"en", "he"}:
        requested_lang = "en"

    word_is_hebrew = _contains_hebrew_letters(raw_word)
    if word_is_hebrew and requested_lang == "he":
        # For Hebrew source words we keep definitions in English to avoid transliteration-heavy round-trips.
        requested_lang = "en"

    if word_is_hebrew:
        meaning, source = _lookup_hebrew_word_meaning(raw_word)
    else:
        meaning, source = _lookup_english_word_meaning(raw_word)

    if meaning and requested_lang == "he" and not word_is_hebrew and not _contains_hebrew_letters(meaning):
        translated_meaning, translated_source = _translate_english_text_online(
            meaning)
        if translated_meaning:
            meaning = translated_meaning
            source_parts = [part for part in [
                source, translated_source] if part]
            source = "+".join(source_parts)

    alternatives = _collect_word_meaning_alternatives(
        raw_word=raw_word,
        primary_meaning=meaning,
        word_is_hebrew=word_is_hebrew,
    )
    if alternatives:
        meaning = alternatives[0]

    if not meaning:
        return jsonify({
            "word": raw_word,
            "meaning": "",
            "alternatives": [],
            "source": "",
            "status": "not_found",
            "lang": requested_lang,
        }), 404

    return jsonify({
        "word": raw_word,
        "meaning": meaning,
        "alternatives": alternatives,
        "source": source,
        "status": "ok",
        "lang": requested_lang,
    })


def _chapter_export_plain_text(title, ref, lines):
    header = [str(title or "").strip(), str(ref or "").strip(), ""]
    body = []
    for idx, line in enumerate(lines, start=1):
        segment = str(line.get("segment") or idx).strip()
        he = str(line.get("he") or "").strip()
        en = str(line.get("en") or "").strip()
        body.append(f"Segment {segment}")
        if he:
            body.append(f"Hebrew: {he}")
        if en:
            body.append(f"English: {en}")
        body.append("")

    return "\n".join(header + body).strip() + "\n"


def _wrap_text_for_export(text, max_chars=96):
    value = re.sub(r"\s+", " ", str(text or "").strip())
    if not value:
        return [""]

    words = value.split(" ")
    chunks = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if len(candidate) <= max_chars:
            current = candidate
            continue
        if current:
            chunks.append(current)
        current = word
    if current:
        chunks.append(current)

    return chunks or [value]


@routes_library.route("/api/export/chapter", methods=["POST"])
@maybe_require_clerk_auth
def export_chapter():
    payload = request.get_json(silent=True) or {}
    title = str(payload.get("title") or payload.get(
        "label") or "shelah-chapter").strip()
    ref = str(payload.get("ref") or "").strip()
    export_format = str(payload.get("format") or "txt").strip().lower()
    lines = payload.get("lines") if isinstance(
        payload.get("lines"), list) else []

    if export_format not in {"txt", "docx", "pdf"}:
        return jsonify({"error": "Unsupported export format"}), 400

    normalized_lines = []
    for idx, line in enumerate(lines, start=1):
        if not isinstance(line, dict):
            continue
        he = str(line.get("he") or "").strip()
        en = str(line.get("en") or "").strip()
        if not he and not en:
            continue
        normalized_lines.append({
            "segment": str(line.get("segment") or idx).strip(),
            "he": he,
            "en": en,
        })

    if not normalized_lines:
        return jsonify({"error": "No chapter lines available to export"}), 400

    file_safe = re.sub(r"[^a-z0-9]+", "-", title.lower()
                       ).strip("-") or "shelah-chapter"
    plain_text = _chapter_export_plain_text(title, ref, normalized_lines)

    if export_format == "txt":
        txt_buffer = io.BytesIO(plain_text.encode("utf-8"))
        txt_buffer.seek(0)
        return send_file(
            txt_buffer,
            as_attachment=True,
            download_name=f"{file_safe}.txt",
            mimetype="text/plain; charset=utf-8",
        )

    if export_format == "docx":
        try:
            from docx import Document as _Document
        except Exception:
            _Document = None
        if _Document is None:
            return jsonify({"error": "DOCX export is unavailable on this server"}), 503

        document = _Document()
        document.add_heading(title or "Sh'elah Chapter", level=1)
        if ref:
            document.add_paragraph(ref)

        for idx, line in enumerate(normalized_lines, start=1):
            segment = line.get("segment") or idx
            document.add_paragraph(f"Segment {segment}")
            if line.get("he"):
                document.add_paragraph(line["he"])
            if line.get("en"):
                document.add_paragraph(line["en"])

        docx_buffer = io.BytesIO()
        document.save(docx_buffer)
        docx_buffer.seek(0)
        return send_file(
            docx_buffer,
            as_attachment=True,
            download_name=f"{file_safe}.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    try:
        from reportlab.lib.pagesizes import LETTER as _LETTER
        from reportlab.pdfgen import canvas as _canvas
    except Exception:
        _LETTER = None
        _canvas = None
    if _canvas is None or _LETTER is None:
        return jsonify({"error": "PDF export is unavailable on this server"}), 503

    pdf_buffer = io.BytesIO()
    pdf = _canvas.Canvas(pdf_buffer, pagesize=_LETTER)
    width, height = _LETTER
    y = height - 48
    left = 42

    def draw_line(text):
        nonlocal y
        for chunk in _wrap_text_for_export(text, max_chars=96):
            if y < 48:
                pdf.showPage()
                y = height - 48
            pdf.drawString(left, y, chunk)
            y -= 14

    draw_line(title or "Sh'elah Chapter")
    if ref:
        draw_line(ref)
    draw_line("")

    for idx, line in enumerate(normalized_lines, start=1):
        draw_line(f"Segment {line.get('segment') or idx}")
        if line.get("he"):
            draw_line(f"Hebrew: {line['he']}")
        if line.get("en"):
            draw_line(f"English: {line['en']}")
        draw_line("")

    pdf.save()
    pdf_buffer.seek(0)
    return send_file(
        pdf_buffer,
        as_attachment=True,
        download_name=f"{file_safe}.pdf",
        mimetype="application/pdf",
    )


@routes_library.route("/api/library/search")
def library_search():
    """Full-text search across Sefaria texts with report-based removal/fix filtering."""
    from backend.sefaria_library import search_library
    query = request.args.get("q", "")
    size = _coerce_int(request.args.get("size"), 10, min_value=1, max_value=50)
    metadata_filters = _extract_search_metadata_filters()
    if not query:
        return jsonify([])
    results = search_library(
        query, size=size, metadata_filters=metadata_filters)
    return jsonify(results)


@routes_library.route("/api/search/suggest")
def search_suggest():
    """Omnibox suggestions: texts, prayers, communities, and AI query option."""
    from backend.sefaria_library import search_library, get_liturgy_books

    query = (request.args.get("q", "") or "").strip()
    size = _coerce_int(request.args.get("size"), 8, min_value=1, max_value=20)
    metadata_filters = _extract_search_metadata_filters()
    if not query:
        return jsonify([])

    q_lower = query.lower()
    suggestions = []
    seen = set()

    def add_item(item_type, label, value, subtitle="", score=0, label_he="", subtitle_he=""):
        key = (item_type, (value or "").lower())
        if key in seen:
            return
        seen.add(key)
        suggestions.append({
            "type": item_type,
            "label": label,
            "label_he": label_he,
            "value": value,
            "subtitle": subtitle,
            "subtitle_he": subtitle_he,
            "score": score,
        })

    alias_ref = QUICK_TEXT_ALIASES.get(q_lower)
    if alias_ref:
        add_item("text", alias_ref, alias_ref, "Popular Torah alias", 100)

    for community in COMMUNITIES.keys():
        if q_lower in community.lower():
            add_item("community", community, community,
                     "Community customs", 90)

    for alias, canonical in COMMUNITY_ALIASES.items():
        if q_lower in alias and canonical in COMMUNITIES:
            add_item("community", canonical, canonical,
                     f"Community customs (matched '{alias}')", 88)

    for book in get_liturgy_books(max_items=120):
        title = book.get("title", "")
        if title and q_lower in title.lower():
            add_item("prayer", title, title, "Sefaria liturgy", 85)

    for hit in search_library(query, size=size, metadata_filters=metadata_filters):
        ref = hit.get("ref", "")
        he_ref = (hit.get("heRef", "") or "").strip()
        categories = " > ".join(hit.get("categories", [])[:3])
        if ref:
            add_item(
                "text",
                ref,
                ref,
                categories or "Sefaria text",
                70,
                label_he=he_ref or ref,
            )

    add_item("ask", f"Ask Sh'elah: {query}", query,
             "AI synthesis", 40)

    suggestions.sort(key=lambda x: x.get("score", 0), reverse=True)
    return jsonify(suggestions[:size])


@routes_library.route("/api/text/<path:ref>/links")
def get_text_links(ref):
    """Returns all linked commentaries & parallel texts for a given ref."""
    from backend.sefaria_library import get_linked_texts
    decoded_ref = _decode_route_ref(ref)
    return jsonify(get_linked_texts(decoded_ref))


@routes_library.route("/api/text/<path:ref>/graph")
def get_text_graph(ref):
    """Build a lightweight source graph around a text reference."""
    from backend.sefaria_library import get_linked_texts

    decoded_ref = _decode_route_ref(ref)
    links = get_linked_texts(decoded_ref)
    nodes = [{"id": decoded_ref, "label": decoded_ref, "kind": "root"}]
    edges = []
    seen = {decoded_ref}

    for category, items in (links or {}).items():
        for item in (items or [])[:14]:
            target = (item or {}).get("ref", "")
            if not target:
                continue
            if target not in seen:
                seen.add(target)
                nodes.append({
                    "id": target,
                    "label": target,
                    "kind": "linked",
                    "category": category,
                })
            edges.append({
                "source": decoded_ref,
                "target": target,
                "label": category,
            })

    return jsonify({
        "ref": ref,
        "nodes": nodes,
        "edges": edges,
    })


@routes_library.route("/api/library/category/<path:category>")
def library_category(category):
    """Returns all books in a given Sefaria category."""
    from backend.sefaria_library import get_category_contents
    return jsonify(get_category_contents(category))


# ─── TEXTS INDEX (for top menu) ───────────────────────────────────────────────
@routes_library.route("/api/texts-index")
def get_texts_index():
    """Returns complete index of browsable texts: prayers, communities, Sefaria."""
    from backend.sefaria_library import get_liturgy_books

    return jsonify({
        "siddur": {
            "title": "Sefaria Prayer Books",
            "items": [b.get("title") for b in get_liturgy_books(max_items=200)]
        },
        "merkava": {
            "title": "Community Customs (Merkava)",
            "items": list(COMMUNITIES.keys())
        },
        "sefaria": {
            "title": "Jewish Text Library",
            "items": ["Tanakh", "Mishnah", "Talmud", "Halakhah", "Kabbalah"]
        }
    })
