"""
Main Flask application for Sh'elah.

What this file owns:
- App bootstrapping, environment wiring, and cache/session policy.
- Public web routes (/, manifest, service worker) and all JSON API routes.
- Integration glue for Supabase preferences, Clerk auth checks, and Sefaria-backed text/prayer/community APIs.
- Calendar and zmanim delivery used by the dashboard (including Hebcal-backed holiday/parasha endpoints).

How to navigate this file:
1) Configuration and helper utilities near the top.
2) Auth and Supabase client helpers.
3) Route handlers grouped by feature (health/devtools, preferences, library/text, prayers, communities, calendar/zmanim).
"""

import json
import re
import io
from concurrent.futures import ThreadPoolExecutor
import requests
from flask import Flask, render_template, request, jsonify, session, g, send_from_directory, send_file
from dotenv import load_dotenv
import time
import os
from typing import Any
from datetime import date as greg_date, timedelta, datetime
from functools import wraps
from urllib.parse import quote, unquote
from pathlib import Path
from uuid import uuid4

import jwt
try:
    from flask_limiter import Limiter
except Exception:
    Limiter = None

try:
    from supabase import create_client
    try:
        from supabase.lib.client_options import SyncClientOptions
    except Exception:
        SyncClientOptions = None
except Exception:
    create_client = None
    SyncClientOptions = None

from pyluach import dates as pyluach_dates

from backend.data_service import ShelahEngine
from backend import sefaria
from backend import claude
from backend import search
from backend.logging_setup import setup_logging
from backend.health_check import health as api_health

try:
    from docx import Document
except Exception:
    Document = None

try:
    from reportlab.lib.pagesizes import LETTER
    from reportlab.pdfgen import canvas
except Exception:
    LETTER = None
    canvas = None

# Maps each prayer name to its constituent Sefaria "Siddur Sefard" refs for full text
SIDDUR_SECTION_MAP = {
    "Upon Arising": [
        "Siddur Sefard, Upon Arising, Modeh Ani",
        "Siddur Sefard, Upon Arising, Introductory Prayers",
        "Siddur Sefard, Upon Arising, Upon Entering Synagogue",
    ],
    "Weekday Shacharit": [
        "Siddur Sefard, Weekday Shacharit, Morning Blessings",
        "Siddur Sefard, Weekday Shacharit, Blessings on Torah",
        "Siddur Sefard, Weekday Shacharit, Morning Prayer",
        "Siddur Sefard, Weekday Shacharit, The Shema",
        "Siddur Sefard, Weekday Shacharit, Amidah",
        "Siddur Sefard, Weekday Shacharit, Tachanun",
        "Siddur Sefard, Weekday Shacharit, Aleinu",
    ],
    "Weekday Mincha": [
        "Siddur Sefard, Weekday Mincha, Amidah",
        "Siddur Sefard, Weekday Mincha, Tachanun",
    ],
    "Weekday Maariv": [
        "Siddur Sefard, Weekday Maariv, The Shema",
        "Siddur Sefard, Weekday Maariv, Amidah",
    ],
    "Shabbat Shacharit": [
        "Siddur Sefard, Shabbat Morning Services, Pesukei D'Zimrah",
        "Siddur Sefard, Shabbat Morning Services, Amidah",
        "Siddur Sefard, Shabbat Morning Services, Shabbat Torah Reading",
    ],
    "Shabbat Mincha": [
        "Siddur Sefard, Shabbat Mincha, Amidah",
    ],
    "Kiddush": [
        "Siddur Sefard, Shabbat Evening Meal, Shabbat Eve Kiddush",
        "Siddur Sefard, Shabbat Day Meal, Shabbat Day Kiddush",
    ],
    "Havdalah": [
        "Siddur Sefard, Motzaei Shabbat , Havdala",
    ],
    "Bedtime Shema": [
        "Siddur Sefard, Bedtime Shema",
    ],
    "Kiddush Levanah": [
        "Siddur Sefard, Kiddush Levanah",
    ],
    "Holiday Prayers": [
        "Siddur Sefard, Holidays, Yom Tov Eve Kiddush",
        "Siddur Sefard, Holidays, Yizkor",
        "Siddur Sefard, Rosh Chodesh, Hallel",
    ],
}

ANSWER_MODES = {"balanced", "practical", "sources", "strict"}

DEVTOOLS_STATS = {
    "answers_total": 0,
    "fallback_answers": 0,
    "strict_blocks": 0,
    "segment_reports": 0,
}

# Bounded in-memory caches — evict oldest entry (insertion-order) when full.
_CACHE_MAX_SIZE = 512


def _bounded_cache_set(cache: dict, key, value, maxsize: int = _CACHE_MAX_SIZE) -> None:
    """Insert into a plain dict while capping its size by evicting the oldest key."""
    if key not in cache and len(cache) >= maxsize:
        cache.pop(next(iter(cache)), None)
    cache[key] = value


ASK_RESPONSE_CACHE: dict = {}
ASK_RESPONSE_CACHE_TTL_SECONDS = 90

QUICK_TEXT_ALIASES = {
    "genesis": "Genesis 1",
    "bereishit": "Genesis 1",
    "exodus": "Exodus 1",
    "shemot": "Exodus 1",
    "leviticus": "Leviticus 1",
    "vayikra": "Leviticus 1",
    "numbers": "Numbers 1",
    "bamidbar": "Numbers 1",
    "deuteronomy": "Deuteronomy 1",
    "devarim": "Deuteronomy 1",
    "psalms": "Psalms 1",
    "tehillim": "Psalms 1",
    "proverbs": "Proverbs 1",
    "mishlei": "Proverbs 1",
    "jonathan sacks": "Covenant and Conversation; Genesis; The Book of the Beginnings, Living with the Times; The Parasha",
    "jonathan sacks essays": "The Jonathan Sacks Haggadah; Essays, The Story of Stories",
    "jonathan sacks haggadah essays": "The Jonathan Sacks Haggadah; Essays, The Story of Stories",
    "covenant and conversation": "Covenant and Conversation; Genesis; The Book of the Beginnings, Living with the Times; The Parasha",
    "everett fox": "The Early Prophets, by Everett Fox, Joshua, Part I; Preparations for Conquest",
    "the early prophets by everett fox": "The Early Prophets, by Everett Fox, Joshua, Part I; Preparations for Conquest",
    "the early prophets, by everett fox": "The Early Prophets, by Everett Fox, Joshua, Part I; Preparations for Conquest",
    "the five books of moses by everett fox": "The Five Books of Moses, by Everett Fox, Translator's Preface",
    "the five books of moses, by everett fox": "The Five Books of Moses, by Everett Fox, Translator's Preface",
}

HEBREW_DIACRITICS_RE = re.compile(r"[\u0591-\u05C7]")
HEBREW_LETTER_RE = re.compile(r"[\u05D0-\u05EA]")

TRANSLATION_CACHE: dict = {}
TRANSLATION_SOURCE_CACHE: dict = {}

GOOGLE_TRANSLATE_API_URL = "https://translate.googleapis.com/translate_a/single"
MYMEMORY_TRANSLATE_API_URL = "https://api.mymemory.translated.net/get"

HEBREW_WORD_GLOSSARY = {
    "שבת": "Shabbat, the seventh day of rest.",
    "תורה": "Torah, the Five Books of Moses and Torah teaching.",
    "תפילה": "Prayer.",
    "מצוה": "Mitzvah, a divine commandment.",
    "מצווה": "Mitzvah, a divine commandment.",
    "הלכה": "Halakhah, practical Jewish law.",
    "מנהג": "Minhag, accepted communal custom.",
    "תשובה": "Teshuvah, repentance and return.",
    "ברכה": "Berakhah, blessing.",
    "פסח": "Pesach, the festival of the Exodus.",
    "סוכות": "Sukkot, the festival of booths.",
    "שבועות": "Shavuot, festival marking Matan Torah.",
    "ראש": "Head or beginning.",
    "שלום": "Peace, well-being, or greeting.",
    "חסד": "Kindness or loving-kindness.",
    "אמת": "Truth.",
    "יראה": "Awe or reverence.",
    "אהבה": "Love.",
}

HEBREW_INTERPRETIVE_GLOSSARY = {
    "ברא": ["create", "fashion", "bring into being"],
    "עשה": ["make", "do", "perform"],
    "אמר": ["say", "speak", "declare"],
    "הלך": ["go", "walk", "proceed"],
    "שמר": ["guard", "keep", "observe"],
}


def _get_cached_ask_payload(cache_key):
    entry = ASK_RESPONSE_CACHE.get(cache_key)
    if not entry:
        return None
    ts = float(entry.get("ts") or 0.0)
    if (time.time() - ts) > ASK_RESPONSE_CACHE_TTL_SECONDS:
        ASK_RESPONSE_CACHE.pop(cache_key, None)
        return None
    payload = entry.get("payload")
    if not isinstance(payload, dict):
        return None
    try:
        # Return a detached copy so callers can mutate response payload safely.
        return json.loads(json.dumps(payload))
    except Exception:
        return None


def _set_cached_ask_payload(cache_key, payload):
    if not isinstance(payload, dict):
        return
    try:
        _bounded_cache_set(ASK_RESPONSE_CACHE, cache_key, {
            "ts": time.time(),
            "payload": json.loads(json.dumps(payload)),
        })
    except Exception:
        return


APP_ROOT = Path(__file__).resolve().parent
SEFARIA_SEARCH_WRAPPER_URL = "https://www.sefaria.org.il/api/search-wrapper"

HALAKHIC_CORPUS_ALIASES = {
    "Shulchan Arukh": [
        "shulchan arukh",
        "shulchan aruch",
        "orach chayim",
        "yoreh de'ah",
        "yoreh deah",
        "even haezer",
        "choshen mishpat",
    ],
    "Rambam": [
        "rambam",
        "mishneh torah",
        "moses maimonides",
    ],
    "Mishnah Berurah": [
        "mishnah berurah",
        "mishna berura",
    ],
    "Talmud": [
        "talmud",
        "bavli",
        "yerushalmi",
    ],
    "Gemara": [
        "gemara",
        "talmud",
        "tractate",
    ],
}

QUERY_STOPWORDS = {
    "the", "and", "for", "with", "from", "that", "this", "are", "was", "were", "have", "has",
    "during", "about", "into", "when", "what", "where", "which", "does", "is", "can", "may", "if",
    "allowed", "halacha", "halakhah", "question", "please", "tell", "me", "us", "you",
    "איך", "מה", "האם", "עם", "של", "על", "גם", "לא", "כן",
}

HEBREW_PREFIXES = ("ו", "ה", "ל", "ב", "ש", "מ")

WEB_LAST_RESORT_WARNING = "⚠️ **WARNING:** No matches found in Sefaria or verified customs. The following info is from the general web and may not be Halakhically accurate. Consult a Rabbi."
WEB_LAST_RESORT_WARNING_PLAIN = WEB_LAST_RESORT_WARNING.replace("**", "")
RABBI_FINAL_RULING_FOOTER = "Please consult with your local Rabbi for a final ruling."
INTERNAL_AI_KNOWLEDGE_DISCLAIMER = (
    "Note: This information was derived from General Halakhic Knowledge "
    f"as the specific database source was unavailable. {RABBI_FINAL_RULING_FOOTER}"
)
WEB_FALLBACK_TRUST_TERMS = {
    "halach", "halakh", "jewish", "judaism", "torah", "talmud", "shabbat",
    "yom tov", "kashrut", "tefillin", "mezuzah", "sefaria", "hebrewbooks",
    "peninei", "yeshivat har bracha", "yhb", "zmanim", "hebrew date", "calendar",
}
WEB_FALLBACK_BLOCKLIST_TERMS = {
    "biblegateway", "biblehub", "biblestudytools", "king james", "new testament",
    "gospel", "church", "jesus", "christian bible",
}

CLOCK_TIME_LATEX_RE = re.compile(
    r"\$(\d{1,2}:\d{2})\s*\\{1,2}text\{\s*(AM|PM)\s*\}\$",
    re.IGNORECASE,
)
DEBUG_OUTPUT_LINE_PATTERNS = [
    re.compile(r"^\s*#{0,6}\s*conflict\s*flag[s]?\b.*$", re.IGNORECASE),
    re.compile(r"^\s*[-*]\s*conflict\s*flag[s]?\b.*$", re.IGNORECASE),
    re.compile(
        r"^\s*(?:[-*]\s*)?source\s*:\s*community\s*knowledge\b.*$", re.IGNORECASE),
    re.compile(
        r"^\s*(?:[-*]\s*)?no\s+primary\s+sefaria\s+snippet\b.*$", re.IGNORECASE),
]
SECTION_KEY_VALUE_RE = re.compile(
    r"^\s*(?P<key>[A-Za-z][A-Za-z0-9 /()'&-]{2,40}):\s*(?P<value>.+)?$"
)
BOLD_HEADER_RE = re.compile(
    r"^\s*\*\*(?P<title>[A-Za-z][A-Za-z0-9 /()'&-]{2,40})\*\*:?\s*(?P<rest>.*)$"
)
HALAKHIC_VERDICT_RE = re.compile(
    r"\b(prohibited|forbidden|permitted|required|obligatory|invalid|valid|asur|assur|mutar)\b",
    re.IGNORECASE,
)
DOMAIN_REFUSAL_MESSAGE_RE = re.compile(
    r"^Sh'elah is a specialized tool for Halakhic and communal knowledge\. "
    r"I cannot assist with .+?, as it falls outside my specialized domain\."
    r"(?:\s+Please consult with your local Rabbi for a final ruling\.)?$",
    re.DOTALL,
)
UI_SECTION_KEYS = {
    "ruling",
    "reason",
    "conditions",
    "exceptions",
    "practical steps",
    "practical guidance",
    "sources",
    "summary",
}
HALAKHIC_VERDICT_LABELS = {
    "prohibited": "Prohibited",
    "forbidden": "Forbidden",
    "permitted": "Permitted",
    "required": "Required",
    "obligatory": "Obligatory",
    "invalid": "Invalid",
    "valid": "Valid",
    "asur": "Asur",
    "assur": "Assur",
    "mutar": "Mutar",
}

# Topic-first direct source anchors for high-signal chapter targeting.
DIRECT_TOPIC_SOURCE_MAP = {
    "omer": {
        "triggers": ["omer", "sefira", "sefirah", "sefirat haomer", "lag baomer", "lag ba'omer", "haircut", "haircuts"],
        "citations": [
            "Shulchan Arukh, Orach Chayim 489",
            "Shulchan Arukh, Orach Chayim 493",
        ],
        "broad_terms": ["omer", "sefirat haomer", "sefira", "haircuts", "mourning customs"],
    },
    "shabbat": {
        "triggers": ["shabbat", "shabbos", "melacha", "havdalah", "kiddush"],
        "citations": [
            "Shulchan Arukh, Orach Chayim 242",
            "Shulchan Arukh, Orach Chayim 318",
        ],
        "broad_terms": ["shabbat", "melacha", "havdalah", "kiddush", "nightfall"],
    },
    "kashrut": {
        "triggers": ["kashrut", "kosher", "basar", "chalav", "meat", "dairy", "treif", "treife"],
        "citations": [
            "Shulchan Arukh, Yoreh De'ah 87",
            "Shulchan Arukh, Yoreh De'ah 89",
        ],
        "broad_terms": ["kashrut", "kosher", "meat and milk", "basar bechalav", "yoreh deah"],
    },
    "niddah": {
        "triggers": ["niddah", "nidda", "mikveh", "taharah", "family purity"],
        "citations": [
            "Shulchan Arukh, Yoreh De'ah 183",
            "Shulchan Arukh, Yoreh De'ah 197",
        ],
        "broad_terms": ["niddah", "family purity", "mikveh", "taharah", "yoreh deah"],
    },
}

QUERY_BROADENER_MAP = {
    "omer": ["sefirat haomer", "sefira", "haircuts"],
    "sefirah": ["omer", "sefirat haomer"],
    "sefira": ["omer", "sefirat haomer"],
    "haircuts": ["haircut", "mourning customs", "omer"],
    "shabbos": ["shabbat", "melacha", "havdalah"],
    "shabbat": ["melacha", "kiddush", "havdalah"],
    "kashrut": ["kosher", "yoreh deah", "meat and milk"],
    "kosher": ["kashrut", "yoreh deah", "meat and milk"],
    "niddah": ["family purity", "mikveh", "taharah"],
}


def _strip_model_web_warning_prefix(answer_text):
    text = str(answer_text or "").strip()
    if not text:
        return ""

    for marker in (WEB_LAST_RESORT_WARNING, WEB_LAST_RESORT_WARNING_PLAIN):
        if text.startswith(marker):
            text = text[len(marker):].lstrip()

    # Guard against model-inserted markdown separators glued to the warning line.
    text = re.sub(r"^-{3,}\s*", "", text)
    return text.strip()


def _join_with_and(parts):
    values = [str(p or "").strip() for p in parts if str(p or "").strip()]
    if not values:
        return ""
    if len(values) == 1:
        return values[0]
    if len(values) == 2:
        return f"{values[0]} and {values[1]}"
    return f"{', '.join(values[:-1])}, and {values[-1]}"


def _build_source_attribution_note(*, has_sefaria=False, has_customs=False, has_whitelisted_external=False, has_general_web=False, has_internal_knowledge=False):
    if has_internal_knowledge or not any((has_sefaria, has_customs, has_whitelisted_external, has_general_web)):
        return INTERNAL_AI_KNOWLEDGE_DISCLAIMER

    sources = []
    if has_sefaria:
        sources.append("Sefaria")
    if has_customs:
        sources.append("Community Customs")
    if has_whitelisted_external:
        sources.append("Halachipedia / HebrewBooks / YHB")
    if has_general_web:
        sources.append("General Web Context")

    joined_sources = _join_with_and(sources)
    return (
        f"Note: This information was pulled from {joined_sources}. "
        f"{RABBI_FINAL_RULING_FOOTER}"
    )


def _compose_answer_with_prefixes(body_text, *, include_web_warning=False, source_attribution_note=""):
    body = str(body_text or "").strip()
    if not body:
        return ""

    blocks = []
    if include_web_warning:
        blocks.append(WEB_LAST_RESORT_WARNING)

    attribution = str(source_attribution_note or "").strip()
    if attribution:
        blocks.append(attribution)

    if blocks:
        blocks.append(body)
        return "\n\n".join(blocks)

    return body


def _strip_source_attribution_prefix(answer_text):
    text = str(answer_text or "").strip()
    if not text:
        return ""

    lower_text = text.lower()
    if not lower_text.startswith("note: this information was"):
        return text

    # Preferred shape is: note line, blank line, then body.
    if "\n\n" in text:
        return text.split("\n\n", 1)[1].lstrip()

    # If model emitted note + footer in one line, remove up to the footer and keep remainder.
    footer_idx = text.find(RABBI_FINAL_RULING_FOOTER)
    if footer_idx != -1:
        remainder = text[footer_idx + len(RABBI_FINAL_RULING_FOOTER):]
        return remainder.lstrip(" \n\t:-")

    # Otherwise only strip the first line and keep any remaining lines.
    first_newline_idx = text.find("\n")
    if first_newline_idx != -1:
        return text[first_newline_idx + 1:].lstrip()

    return text


def _normalize_ai_answer(answer_text, include_web_warning=False, source_attribution_note="", allow_empty_fallback=True):
    body = _strip_model_web_warning_prefix(answer_text)
    body = _strip_source_attribution_prefix(body)
    body = CLOCK_TIME_LATEX_RE.sub(
        lambda m: f"{m.group(1)} {m.group(2).upper()}", body)

    # Preserve the domain guardrail refusal message exactly as emitted.
    if DOMAIN_REFUSAL_MESSAGE_RE.fullmatch(body):
        if RABBI_FINAL_RULING_FOOTER not in body:
            return f"{body}\n\n{RABBI_FINAL_RULING_FOOTER}"
        return body

    if not body:
        if allow_empty_fallback:
            body = "No verified source found"
        else:
            return ""

    prefix_blocks = []
    if include_web_warning:
        prefix_blocks.append(WEB_LAST_RESORT_WARNING)

    attribution = str(source_attribution_note or "").strip()
    if attribution:
        prefix_blocks.append(attribution)
    elif body.lower() != "no verified source found" and RABBI_FINAL_RULING_FOOTER not in body:
        prefix_blocks.append(RABBI_FINAL_RULING_FOOTER)

    if prefix_blocks:
        return "\n\n".join(prefix_blocks + [body])

    return body


def _should_drop_debug_line(line_text):
    return any(pattern.match(line_text) for pattern in DEBUG_OUTPUT_LINE_PATTERNS)


def _bold_halakhic_verdicts(text):
    def _replace(match):
        token = str(match.group(0) or "")
        canonical = HALAKHIC_VERDICT_LABELS.get(token.lower(), token)
        return f"**{canonical}**"

    return HALAKHIC_VERDICT_RE.sub(_replace, text)


def _normalize_answer_line(raw_line):
    line = str(raw_line or "").strip()
    if not line:
        return [""]

    if _should_drop_debug_line(line):
        return []

    line = re.sub(r"^\s*[•*]\s+", "- ", line)
    line = re.sub(r"^\s*[-–]\s+", "- ", line)

    if line.startswith("# "):
        line = f"## {line[2:].strip()}"

    bold_header = BOLD_HEADER_RE.match(line)
    if bold_header:
        title = bold_header.group("title").strip().rstrip(":")
        rest = (bold_header.group("rest") or "").strip()
        output = [f"### {title}"]
        if rest:
            output.append(_bold_halakhic_verdicts(rest))
        return output

    key_value = SECTION_KEY_VALUE_RE.match(line)
    if key_value:
        key = key_value.group("key").strip()
        value = (key_value.group("value") or "").strip()
        if key.lower() in UI_SECTION_KEYS:
            output = [f"### {key}"]
            if value:
                output.append(_bold_halakhic_verdicts(value))
            return output

    return [_bold_halakhic_verdicts(line)]


def _collapse_markdown_spacing(lines):
    normalized = []
    prev_blank = True

    for line in lines:
        text = str(line or "")
        is_blank = not text.strip()

        if is_blank:
            if not prev_blank:
                normalized.append("")
            prev_blank = True
            continue

        if text.startswith("##") or text.startswith("###"):
            if normalized and normalized[-1] != "":
                normalized.append("")
            normalized.append(text)
            prev_blank = False
            continue

        normalized.append(text)
        prev_blank = False

    while normalized and not normalized[0].strip():
        normalized.pop(0)
    while normalized and not normalized[-1].strip():
        normalized.pop()

    return normalized


def _format_ui_answer(answer_text):
    lines = []
    for raw_line in str(answer_text or "").splitlines():
        lines.extend(_normalize_answer_line(raw_line))

    lines = _collapse_markdown_spacing(lines)
    if not lines:
        return ""

    has_headers = any(line.startswith("##") or line.startswith("###")
                      for line in lines)
    if not has_headers and str(answer_text or "").strip().lower() != "no verified source found":
        lines = ["## Ruling", "", lines[0]] + lines[1:]
        lines = _collapse_markdown_spacing(lines)

    return "\n".join(lines).strip()


def _strip_common_hebrew_prefixes(token):
    value = str(token or "").strip()
    if not value:
        return value
    if not HEBREW_LETTER_RE.search(value):
        return value

    normalized = value
    # Allow stacked prefixes but keep a meaningful stem length.
    while len(normalized) > 2 and normalized[0] in HEBREW_PREFIXES:
        normalized = normalized[1:]

    return normalized or value


def _expand_hebrew_keyword_forms(token):
    base = str(token or "").strip().lower()
    if not base:
        return []

    expanded = [base]
    stripped = _strip_common_hebrew_prefixes(base)
    if stripped and stripped not in expanded:
        expanded.append(stripped)

    return expanded


def _extract_query_keywords(query, max_keywords=8):
    tokens = re.findall(r"[A-Za-z\u0590-\u05FF]{3,}", str(query or "").lower())
    keywords = []
    for token in tokens:
        for normalized_token in _expand_hebrew_keyword_forms(token):
            if normalized_token in QUERY_STOPWORDS:
                continue
            if normalized_token not in keywords:
                keywords.append(normalized_token)
            if len(keywords) >= max_keywords:
                break
        if len(keywords) >= max_keywords:
            break
    return keywords


def _query_search_wrapper(query_text, size=12):
    payload = {
        "type": "text",
        "query": query_text,
        "field": "naive_lemmatizer",
        "source_proj": True,
        "slop": 10,
        "start": 0,
        "size": size,
        "filters": [],
        "filter_fields": [],
        "aggs": [],
        "sort_method": "score",
        "sort_fields": ["pagesheetrank"],
        "sort_reverse": False,
        "sort_score_missing": 0.04,
    }

    try:
        resp = requests.post(
            SEFARIA_SEARCH_WRAPPER_URL,
            json=payload,
            timeout=8,
        )
        if not resp.ok:
            return []
        data = resp.json() if resp.content else {}
        return ((data.get("hits") or {}).get("hits") or [])
    except Exception:
        return []


def _match_corpus(hit_source, aliases):
    if not isinstance(hit_source, dict):
        return False

    categories = hit_source.get("categories", [])
    title_variants = hit_source.get("titleVariants", [])
    haystack = " ".join([
        str(hit_source.get("ref") or ""),
        str(hit_source.get("path") or ""),
        " ".join(categories if isinstance(categories, list) else []),
        " ".join(title_variants if isinstance(title_variants, list) else []),
    ]).lower().replace("_", " ")

    return any(alias in haystack for alias in aliases)


def _extract_hit_snippet(hit_source):
    for key in ("naive_lemmatizer", "exact", "content"):
        raw = hit_source.get(key, "")
        if isinstance(raw, str) and raw.strip():
            return re.sub(r"\s+", " ", raw).strip()[:340]
    return ""


def _dedupe_ordered_text(values, max_items=None):
    collected = []
    seen = set()

    for value in values or []:
        normalized = re.sub(r"\s+", " ", str(value or "")).strip()
        if not normalized:
            continue

        key = normalized.lower()
        if key in seen:
            continue

        seen.add(key)
        collected.append(normalized)

        if max_items and len(collected) >= max_items:
            break

    return collected


def _match_direct_topics(question, keywords):
    haystack = f"{question} {' '.join(keywords)}".lower()
    matched = []

    for topic, config in DIRECT_TOPIC_SOURCE_MAP.items():
        triggers = config.get("triggers", [])
        if any(trigger in haystack for trigger in triggers):
            matched.append(topic)

    return matched


def _build_discovery_queries(question, keywords):
    matched_topics = _match_direct_topics(question, keywords)

    specific_queries = []
    broad_terms = list(keywords)

    for topic in matched_topics:
        config = DIRECT_TOPIC_SOURCE_MAP.get(topic, {})
        specific_queries.extend(config.get("citations", []))
        specific_queries.append(f"{topic} shulchan arukh")
        broad_terms.extend(config.get("broad_terms", []))

    for keyword in keywords:
        broad_terms.extend(QUERY_BROADENER_MAP.get(keyword.lower(), []))

    if question:
        specific_queries.append(question)
        broad_terms.extend(_extract_query_keywords(question, max_keywords=12))

    specific_queries = _dedupe_ordered_text(specific_queries, max_items=14)
    broad_terms = _dedupe_ordered_text(broad_terms, max_items=18)

    broad_queries = []
    if broad_terms:
        broad_queries.append(" ".join(broad_terms[:5]))
        if len(broad_terms) >= 8:
            broad_queries.append(" ".join(broad_terms[3:8]))
        broad_queries.extend(broad_terms[:10])
    if question:
        broad_queries.append(question)
    broad_queries = _dedupe_ordered_text(broad_queries, max_items=16)

    if not specific_queries and question:
        specific_queries = [question]

    return {
        "topics": matched_topics,
        "specific_queries": specific_queries,
        "broad_queries": broad_queries,
    }


def _is_sefaria_hit_relevant(hit_source, query_terms):
    if not query_terms:
        return True

    categories = hit_source.get("categories", [])
    title_variants = hit_source.get("titleVariants", [])
    snippet = _extract_hit_snippet(hit_source)
    haystack = " ".join([
        str(hit_source.get("ref") or ""),
        str(hit_source.get("path") or ""),
        " ".join(categories if isinstance(categories, list) else []),
        " ".join(title_variants if isinstance(title_variants, list) else []),
        snippet,
    ]).lower().replace("_", " ")

    terms = []
    for term in query_terms:
        normalized = str(term or "").strip().lower()
        if len(normalized) < 3 or normalized in QUERY_STOPWORDS:
            continue
        terms.append(normalized)

    if not terms:
        return True

    return any(term in haystack for term in terms)


def _collect_global_sefaria_sources(queries, fallback_terms, discovery_stage, priority, max_results=10):
    sources = []
    seen_refs = set()

    per_query_limit = 3 if discovery_stage == "specific-api" else 2

    for query_text in queries:
        normalized_query = str(query_text or "").strip()
        if not normalized_query:
            continue

        hits = _query_search_wrapper(normalized_query, size=80)
        query_terms = _extract_query_keywords(
            normalized_query) or fallback_terms

        added_for_query = 0
        for hit in hits:
            hit_source = hit.get("_source", {}) if isinstance(
                hit, dict) else {}
            if not isinstance(hit_source, dict):
                continue

            if not _is_sefaria_hit_relevant(hit_source, query_terms):
                continue

            ref = str(hit_source.get("ref") or "").strip()
            if not ref or ref in seen_refs:
                continue

            seen_refs.add(ref)
            he_ref = str(hit_source.get("heRef") or "").strip()
            path = str(hit_source.get("path") or "").strip()
            snippet = _extract_hit_snippet(hit_source)

            sources.append({
                "ref": ref,
                "title": ref,
                "lines": [{"en": snippet or f"Matched via global search: {normalized_query}", "he": he_ref}],
                "domain": "Sefaria",
                "corpus": "sefaria-global-search",
                "path": path,
                "priority": priority,
                "status": "fallback",
                "discovery_stage": discovery_stage,
                "search_query": normalized_query,
                "score": hit.get("_score") if isinstance(hit, dict) else None,
            })

            added_for_query += 1
            if added_for_query >= per_query_limit or len(sources) >= max_results:
                break

        if len(sources) >= max_results:
            break

    return sources


def _collect_external_global_sources(queries, keywords, discovery_stage, priority, max_results=6):
    providers = [
        ("Halachipedia", "halachipedia.com", search.search_halachipedia),
        ("HebrewBooks", "hebrewbooks.org", search.search_hebrewbooks),
    ]

    sources = []
    seen = set()

    for query_text in queries:
        normalized_query = str(query_text or "").strip()
        if not normalized_query:
            continue

        for provider_name, domain, provider_search in providers:
            payload = provider_search(normalized_query)
            if not isinstance(payload, dict):
                continue

            title = str(payload.get("title") or "").strip()
            summary = str(payload.get("summary") or "").strip()
            if provider_name == "Halachipedia":
                title = re.sub(r"^\[Halachipedia\]\s*", "", title).strip()

            if not title and not summary:
                continue

            if not _looks_like_trusted_web_match(provider_name.lower(), title, summary, keywords):
                continue

            dedupe_key = (provider_name.lower(), title.lower())
            if dedupe_key in seen:
                continue
            seen.add(dedupe_key)

            url = str(payload.get("url") or "").strip()
            if not url and provider_name == "Halachipedia" and title:
                slug = quote(title.replace(" ", "_"), safe="")
                url = f"https://halachipedia.com/wiki/{slug}" if slug else "https://halachipedia.com"
            if not url and provider_name == "HebrewBooks":
                url = f"https://www.hebrewbooks.org/search.aspx?st=FT&q={quote(normalized_query, safe='')}"

            sources.append({
                "ref": title[:140] or provider_name,
                "title": title[:160] or provider_name,
                "lines": [{"en": summary[:1000], "he": ""}],
                "domain": domain,
                "corpus": "external-global-search",
                "source_provider": provider_name,
                "url": url,
                "priority": priority,
                "status": "fallback",
                "discovery_stage": discovery_stage,
                "search_query": normalized_query,
            })

            if len(sources) >= max_results:
                return sources

    return sources


def _iter_local_json_matches(payload, keywords, file_name, pointer="root"):
    matches = []

    if isinstance(payload, dict):
        for field in ("Minhag", "minhag", "Title", "title"):
            value = payload.get(field)
            if not isinstance(value, str):
                continue
            lowered = value.lower()
            hit_keywords = [kw for kw in keywords if kw in lowered]
            if not hit_keywords:
                continue
            matches.append({
                "file": file_name,
                "field": field,
                "value": value,
                "match_keywords": hit_keywords,
                "pointer": pointer,
            })

        for key, value in payload.items():
            child_pointer = f"{pointer}.{key}"
            matches.extend(_iter_local_json_matches(
                value, keywords, file_name, child_pointer))

    elif isinstance(payload, list):
        for idx, item in enumerate(payload):
            child_pointer = f"{pointer}[{idx}]"
            matches.extend(_iter_local_json_matches(
                item, keywords, file_name, child_pointer))

    return matches


def _find_local_custom_matches(keywords, max_results=12):
    roots = [
        APP_ROOT / ".github" / "customs",
        APP_ROOT / "customs",
    ]

    collected = []
    seen = set()

    for root in roots:
        if not root.exists() or not root.is_dir():
            continue

        for file_path in sorted(root.glob("*.json")):
            try:
                payload = json.loads(file_path.read_text(encoding="utf-8"))
            except Exception:
                continue

            for match in _iter_local_json_matches(payload, keywords, file_path.name):
                key = (
                    match.get("file", ""),
                    match.get("field", ""),
                    str(match.get("value", "")).lower(),
                    match.get("pointer", ""),
                )
                if key in seen:
                    continue
                seen.add(key)
                collected.append(match)
                if len(collected) >= max_results:
                    return collected

    return collected


def _looks_like_trusted_web_match(provider, title, summary, keywords):
    provider_name = str(provider or "").strip().lower()
    title_text = str(title or "").strip()
    summary_text = str(summary or "").strip()
    if not title_text or not summary_text:
        return False

    haystack = f"{title_text} {summary_text}".lower()
    if any(flag in haystack for flag in WEB_FALLBACK_BLOCKLIST_TERMS):
        return False

    if provider_name in {"halachipedia", "hebrewbooks"}:
        return True

    if any(term in haystack for term in WEB_FALLBACK_TRUST_TERMS):
        return True

    # Require relevance to the query if no explicit trust-term signal is present.
    return any(kw in haystack for kw in keywords)


def _build_last_resort_web_sources(question, keywords, max_results=6):
    query = str(question or "").strip()
    keyword_query = " ".join(keywords[:5]).strip()

    halachipedia_queries = []
    if keyword_query:
        halachipedia_queries.append(keyword_query)
        halachipedia_queries.append(f"halakha {keyword_query}".strip())
    if query:
        halachipedia_queries.append(query)

    wiki_titles = [
        "Peninei Halakha",
        "Yeshivat Har Bracha",
        "HebrewBooks",
    ]
    if keyword_query:
        wiki_titles.append(f"Halakha {keyword_query}".strip())
    if query:
        wiki_titles.append(f"Halakha {query}".strip())

    candidates = []

    for q in halachipedia_queries:
        if not q:
            continue
        payload = search.search_halachipedia(q)
        if not isinstance(payload, dict):
            continue
        title = str(payload.get("title") or "").strip()
        summary = str(payload.get("summary") or "").strip()
        clean_title = re.sub(r"^\[Halachipedia\]\s*", "", title).strip()
        if not _looks_like_trusted_web_match("halachipedia", clean_title, summary, keywords):
            continue
        url_slug = quote(clean_title.replace(" ", "_"),
                         safe="") if clean_title else ""
        candidates.append({
            "provider": "Halachipedia",
            "domain": "halachipedia.com",
            "title": clean_title or title,
            "summary": summary,
            "url": f"https://halachipedia.com/wiki/{url_slug}" if url_slug else "https://halachipedia.com",
        })

    for title_query in wiki_titles:
        if not title_query:
            continue
        payload = search.search_wikipedia(title_query)
        if not isinstance(payload, dict):
            continue
        title = str(payload.get("title") or "").strip()
        summary = str(payload.get("summary") or "").strip()
        if not _looks_like_trusted_web_match("wikipedia", title, summary, keywords):
            continue
        url_slug = quote(title.replace(" ", "_"), safe="") if title else ""
        candidates.append({
            "provider": "Wikipedia",
            "domain": "en.wikipedia.org",
            "title": title,
            "summary": summary,
            "url": f"https://en.wikipedia.org/wiki/{url_slug}" if url_slug else "https://en.wikipedia.org",
        })

    deduped = []
    seen = set()
    for item in candidates:
        key = (
            str(item.get("provider") or "").lower(),
            str(item.get("title") or "").strip().lower(),
        )
        if key in seen:
            continue
        seen.add(key)
        deduped.append(item)
        if len(deduped) >= max_results:
            break

    web_sources = []
    for item in deduped:
        title = str(item.get("title") or "Web Source").strip()
        summary = str(item.get("summary") or "").strip()
        web_sources.append({
            "ref": title[:140],
            "title": title[:160],
            "lines": [{"en": summary[:1000], "he": ""}],
            "domain": item.get("domain"),
            "corpus": "general-web",
            "source_provider": item.get("provider"),
            "url": item.get("url"),
            "priority": 3,
            "status": "fallback-web",
        })

    return web_sources


def get_halakhic_sources(query):
    """Global discovery fallback: specific API -> broad API -> internal AI knowledge."""
    question = claude.sanitize_user_query(query)
    keywords = _extract_query_keywords(question)
    if not keywords and question:
        keywords = [question.lower()]

    discovery = _build_discovery_queries(question, keywords)
    topic_matches = discovery.get("topics", [])
    specific_queries = discovery.get("specific_queries", [])
    broad_queries = discovery.get("broad_queries", [])

    specific_sefaria = _collect_global_sefaria_sources(
        specific_queries,
        fallback_terms=keywords,
        discovery_stage="specific-api",
        priority=1,
        max_results=8,
    )
    specific_external = _collect_external_global_sources(
        specific_queries,
        keywords=keywords,
        discovery_stage="specific-api",
        priority=1,
        max_results=4,
    )
    specific_sources = specific_sefaria + specific_external

    if specific_sources:
        return {
            "status": "fallback",
            "fallback_level": "specific-api",
            "query": question,
            "keywords": keywords,
            "topics": topic_matches,
            "specific_queries": specific_queries,
            "broad_queries": broad_queries,
            "sequence": ["specific-api", "broad-api", "internal-ai-knowledge"],
            "counts": {
                "specific_api": len(specific_sources),
                "broad_api": 0,
                "internal_ai": 0,
                "sefaria": len(specific_sefaria),
                "external": len(specific_external),
            },
            "warning": "",
            "internal_disclaimer": "",
            "source_count": len(specific_sources),
            "sources": specific_sources,
        }

    broad_sefaria = _collect_global_sefaria_sources(
        broad_queries,
        fallback_terms=keywords,
        discovery_stage="broad-api",
        priority=2,
        max_results=10,
    )
    broad_external = _collect_external_global_sources(
        broad_queries,
        keywords=keywords,
        discovery_stage="broad-api",
        priority=2,
        max_results=6,
    )
    broad_sources = broad_sefaria + broad_external

    if broad_sources:
        return {
            "status": "fallback",
            "fallback_level": "broad-api",
            "query": question,
            "keywords": keywords,
            "topics": topic_matches,
            "specific_queries": specific_queries,
            "broad_queries": broad_queries,
            "sequence": ["specific-api", "broad-api", "internal-ai-knowledge"],
            "counts": {
                "specific_api": 0,
                "broad_api": len(broad_sources),
                "internal_ai": 0,
                "sefaria": len(broad_sefaria),
                "external": len(broad_external),
            },
            "warning": "",
            "internal_disclaimer": "",
            "source_count": len(broad_sources),
            "sources": broad_sources,
        }

    return {
        "status": "internal-ai-needed",
        "fallback_level": "internal-ai-knowledge",
        "query": question,
        "keywords": keywords,
        "topics": topic_matches,
        "specific_queries": specific_queries,
        "broad_queries": broad_queries,
        "sequence": ["specific-api", "broad-api", "internal-ai-knowledge"],
        "counts": {
            "specific_api": 0,
            "broad_api": 0,
            "internal_ai": 1,
            "sefaria": 0,
            "external": 0,
        },
        "warning": "",
        "internal_disclaimer": INTERNAL_AI_KNOWLEDGE_DISCLAIMER,
        "source_count": 1,
        "sources": [{
            "ref": "Internal Halakhic Knowledge",
            "title": "Internal Halakhic Knowledge",
            "lines": [{
                "en": "No relevant API snippet was found. The answer may rely on internal Halakhic knowledge with the required disclaimer.",
                "he": "",
            }],
            "domain": "internal-ai",
            "corpus": "internal-knowledge",
            "priority": 3,
            "status": "internal-ai-needed",
        }],
    }


def _build_ask_tool_context(engine):
    context = {
        "route": "/ask",
        "auth_enforced": CLERK_ENFORCE_AUTH,
        "trusted_source_priority": "Sefaria, HebrewBooks, Halachipedia, Peninei Halakha (YHB)",
        "factuality_guardrail": "Reject random/non-halakhic domains and avoid generic English Bible websites when validating web context.",
    }

    try:
        zmanim_payload = engine.get_zmanim("standard")
        if isinstance(zmanim_payload, dict):
            metadata = zmanim_payload.get("metadata", {})
            zmanim = zmanim_payload.get("zmanim", {})

            if isinstance(metadata, dict):
                context["civil_date"] = metadata.get("date")
                context["hebrew_date"] = metadata.get("hebrew_date")
                context["parasha"] = metadata.get("parasha")
                context["holiday"] = metadata.get("holiday")
                context["timezone"] = metadata.get("timezone")

            if isinstance(zmanim, dict):
                snapshot = {}
                for key in (
                    "Dawn (16.1° / 72m)",
                    "Sunrise",
                    "Latest Shema (GRA)",
                    "Plag HaMincha",
                    "Sunset",
                    "Nightfall (3 Stars)",
                ):
                    value = str(zmanim.get(key) or "").strip()
                    if value and value != "N/A":
                        snapshot[key] = value

                if snapshot:
                    context["zmanim_snapshot"] = snapshot
    except Exception:
        # Keep ask flow resilient even if zmanim context is unavailable.
        pass

    return context


def _compact_ai_sources(sources, max_sources=8, max_lines=3, max_chars=280):
    """Trim bulky source payloads to the excerpt shape used by the UI."""
    if not isinstance(sources, list):
        return []

    compacted = []
    for src in sources[:max_sources]:
        if not isinstance(src, dict):
            continue

        ref = str(src.get("ref") or "").strip()
        title = str(src.get("title") or ref).strip()
        raw_lines_obj = src.get("lines")
        raw_lines = raw_lines_obj if isinstance(raw_lines_obj, list) else []

        lines = []
        has_valid_content = False
        for row in raw_lines[:max_lines]:
            if not isinstance(row, dict):
                continue

            en = re.sub(r"\s+", " ", str(row.get("en") or "").strip())
            he = re.sub(r"\s+", " ", str(row.get("he") or "").strip())

            # Skip lines that indicate the source was not found
            if en.startswith("Text not found") or en.startswith("Error"):
                continue

            if len(en) > max_chars:
                en = f"{en[:max_chars].rstrip()}..."
            if len(he) > max_chars:
                he = f"{he[:max_chars].rstrip()}..."

            if en or he:
                has_valid_content = True
            lines.append({"en": en, "he": he})

        # Skip sources with no valid content
        if not has_valid_content and not lines:
            continue

        domain = str(src.get("domain") or "").strip()
        source_provider = str(src.get("source_provider") or "").strip()
        url = str(src.get("url") or "").strip()

        entry: dict = {
            "ref": ref[:220],
            "title": title[:220],
            "lines": lines,
        }
        if domain:
            entry["domain"] = domain
        if source_provider:
            entry["source_provider"] = source_provider
        if url:
            entry["url"] = url

        compacted.append(entry)

    return compacted


def _strip_hebrew_diacritics(text):
    return HEBREW_DIACRITICS_RE.sub("", str(text or ""))


def _contains_hebrew_letters(text):
    return bool(HEBREW_LETTER_RE.search(str(text or "")))


def _normalize_lookup_word(text):
    cleaned = _strip_hebrew_diacritics(text).strip()
    cleaned = re.sub(r"\s+", " ", cleaned)
    return cleaned


def _decode_route_ref(value, max_rounds=3):
    """Decode refs that may arrive pre-encoded or double-encoded from clients/proxies."""
    decoded = str(value or "").strip()
    for _ in range(max_rounds):
        next_value = unquote(decoded).strip()
        if next_value == decoded:
            break
        decoded = next_value
    return decoded


def _is_translation_echo(source_text, translated_text):
    src = re.sub(r"\s+", " ", str(source_text or "").strip()).lower()
    dst = re.sub(r"\s+", " ", str(translated_text or "").strip()).lower()
    return bool(src and dst and src == dst)


def _extract_google_translated_text(payload):
    if not isinstance(payload, list) or not payload:
        return ""
    segments = payload[0]
    if not isinstance(segments, list):
        return ""

    chunks = []
    for segment in segments:
        if isinstance(segment, list) and segment:
            chunk = str(segment[0] or "").strip()
            if chunk:
                chunks.append(chunk)

    return re.sub(r"\s+", " ", "".join(chunks)).strip()


def _translate_text_google(text, source_lang, target_lang):
    value = str(text or "").strip()
    if not value:
        return ""

    try:
        resp = requests.get(
            GOOGLE_TRANSLATE_API_URL,
            params={
                "client": "gtx",
                "sl": str(source_lang or "auto").strip() or "auto",
                "tl": str(target_lang or "en").strip() or "en",
                "dt": "t",
                "q": value,
            },
            headers={"User-Agent": "Mozilla/5.0"},
            timeout=2.5,
        )
        if not resp.ok:
            return ""
        payload = resp.json() if resp.content else []
        translated = _extract_google_translated_text(payload)
        if not translated:
            return ""
        if _is_translation_echo(value, translated):
            return ""
        return translated
    except Exception:
        return ""


def _translate_text_mymemory(text, source_lang, target_lang):
    value = str(text or "").strip()
    if not value:
        return ""

    langpair_source = str(source_lang or "auto").strip() or "auto"
    langpair_target = str(target_lang or "en").strip() or "en"

    try:
        resp = requests.get(
            MYMEMORY_TRANSLATE_API_URL,
            params={"q": value, "langpair": f"{langpair_source}|{langpair_target}"},
            timeout=2.5,
        )
        if not resp.ok:
            return ""
        payload = resp.json() if resp.content else {}
        translated = str((payload.get("responseData") or {}).get(
            "translatedText") or "").strip()
        if not translated:
            return ""
        if _is_translation_echo(value, translated):
            return ""
        return translated
    except Exception:
        return ""


def _translate_hebrew_text_google(text):
    value = str(text or "").strip()
    if not value or not _contains_hebrew_letters(value):
        return ""
    translated = _translate_text_google(value, "he", "en")
    if not translated or _is_translation_echo(value, translated):
        return ""
    return translated


def _translate_hebrew_text_mymemory(text):
    value = str(text or "").strip()
    if not value or not _contains_hebrew_letters(value):
        return ""
    translated = _translate_text_mymemory(value, "he", "en")
    if not translated or _is_translation_echo(value, translated):
        return ""
    return translated


def _translate_hebrew_text_online(text):
    """Best-effort Hebrew->English translation, preferring Google with fallback providers."""
    value = _normalize_lookup_word(text)
    if not value:
        return "", ""
    if not _contains_hebrew_letters(value):
        return "", ""

    cache_key = value[:320]
    if cache_key in TRANSLATION_CACHE:
        return TRANSLATION_CACHE.get(cache_key, ""), TRANSLATION_SOURCE_CACHE.get(cache_key, "")

    translated = _translate_hebrew_text_google(cache_key)
    source = "google-translate"

    if not translated:
        translated = _translate_hebrew_text_mymemory(cache_key)
        source = "mymemory-translate"

    _bounded_cache_set(TRANSLATION_CACHE, cache_key, translated)
    _bounded_cache_set(TRANSLATION_SOURCE_CACHE, cache_key,
                       source if translated else "")

    if not translated:
        return "", ""

    return translated, source


def _translate_english_text_online(text):
    """Best-effort English->Hebrew translation for Hebrew UI word-meaning output."""
    value = _normalize_lookup_word(text)
    if not value:
        return "", ""

    cache_key = f"en-he::{value[:320]}"
    if cache_key in TRANSLATION_CACHE:
        return TRANSLATION_CACHE.get(cache_key, ""), TRANSLATION_SOURCE_CACHE.get(cache_key, "")

    translated = _translate_text_google(value, "en", "he")
    source = "google-translate"

    if not translated or _is_translation_echo(value, translated):
        translated = _translate_text_mymemory(value, "en", "he")
        source = "mymemory-translate"

    if _is_translation_echo(value, translated):
        translated = ""

    _bounded_cache_set(TRANSLATION_CACHE, cache_key, translated)
    _bounded_cache_set(TRANSLATION_SOURCE_CACHE, cache_key,
                       source if translated else "")

    if not translated:
        return "", ""

    return translated, source


def _fill_missing_english_lines(text_payload, max_lines=12, max_runtime_seconds=2.5):
    """Fill missing English lines when Hebrew is available and translation can be generated."""
    if not isinstance(text_payload, dict):
        return text_payload

    lines = text_payload.get("lines", [])
    if not isinstance(lines, list) or not lines:
        return text_payload

    translated_count = 0
    translation_sources = set()
    started_at = time.time()
    for line in lines:
        if translated_count >= max_lines:
            break
        if time.time() - started_at > max_runtime_seconds:
            break
        if not isinstance(line, dict):
            continue

        en_value = str(line.get("en") or "").strip()
        he_value = _normalize_lookup_word(line.get("he") or "")
        if en_value or not he_value:
            continue
        if not _contains_hebrew_letters(he_value):
            continue

        generated, source = _translate_hebrew_text_online(he_value[:320])

        if generated:
            line["en"] = generated
            translated_count += 1
            if source:
                translation_sources.add(source)

    if translated_count:
        provider_label = ", ".join(
            sorted(translation_sources)) if translation_sources else "online-translation"
        text_payload["translation_generated"] = True
        text_payload["translation_generated_count"] = translated_count
        text_payload["translation_source"] = provider_label
        text_payload[
            "translation_note"] = f"Automatic English translation added for missing lines ({provider_label})."
        text_payload["en"] = [
            str(line.get("en", "")).strip()
            for line in lines
            if isinstance(line, dict) and str(line.get("en", "")).strip()
        ]

    return text_payload


def _build_trusted_custom_sources(data):
    """Build a stable source list from trusted halachic authorities in community files."""
    if not isinstance(data, dict):
        return []

    candidates = []

    source_registry = data.get("source_registry", {}) if isinstance(
        data.get("source_registry"), dict) else {}
    candidates.extend(source_registry.get("primary", []) if isinstance(
        source_registry.get("primary"), list) else [])

    authorities = data.get("core_halachic_authorities", {}) if isinstance(
        data.get("core_halachic_authorities"), dict) else {}
    for key in (
        "primary_codes",
        "major_rishonim_base",
        "later_ashkenazi_poskim",
        "later_sephardi_poskim",
        "later_moroccan_poskim",
        "later_turkish_poskim",
    ):
        value = authorities.get(key)
        if isinstance(value, list):
            candidates.extend(value)

    deduped = []
    seen = set()
    for item in candidates:
        label = str(item or "").strip()
        if not label:
            continue
        key = label.lower()
        if key in seen:
            continue
        seen.add(key)
        deduped.append(label)

    return deduped[:6]


def _lookup_english_word_meaning(word):
    clean_word = str(word or "").strip().lower()
    if not clean_word:
        return "", ""

    try:
        resp = requests.get(
            f"https://api.dictionaryapi.dev/api/v2/entries/en/{clean_word}",
            timeout=5,
        )
        if not resp.ok:
            return "", ""
        payload = resp.json()
        if not isinstance(payload, list) or not payload:
            return "", ""

        entry = payload[0] if isinstance(payload[0], dict) else {}
        meanings = entry.get("meanings", []) if isinstance(
            entry.get("meanings"), list) else []
        for meaning in meanings:
            definitions = meaning.get(
                "definitions", []) if isinstance(meaning, dict) else []
            for definition in definitions:
                text = str((definition or {}).get("definition") or "").strip()
                if text:
                    return text, "dictionaryapi.dev"
    except Exception:
        return "", ""

    return "", ""


def _normalize_glossary_meaning(value):
    text = re.sub(r"\s+", " ", str(value or "").strip())
    if not text:
        return ""

    if "," in text:
        head, tail = text.split(",", 1)
        if re.fullmatch(r"[A-Za-z'\-\s]{2,40}", head.strip()) and tail.strip():
            return tail.strip()

    return text


def _looks_like_transliteration(text):
    value = re.sub(r"\s+", " ", str(text or "").strip())
    if not value:
        return False
    if not re.fullmatch(r"[A-Za-z'\-\s]{2,80}", value):
        return False

    lower = value.lower()
    tokens = [part for part in lower.split(" ") if part]
    if not tokens:
        return False

    # Clear transliteration signals (e.g. bereshit, elohim, shabbat, tzedakah).
    translit_markers = ("sh", "kh", "tz", "ts", "aa", "ee", "oo", "iy", "ui")
    translit_suffixes = ("im", "ot", "ah", "eh", "it", "ut", "iyyah")

    if any("'" in token or "-" in token for token in tokens):
        return True

    if len(tokens) <= 3 and any(marker in lower for marker in translit_markers):
        return True

    if len(tokens) <= 2 and all(any(token.endswith(suffix) for suffix in translit_suffixes) for token in tokens):
        return True

    if len(tokens) == 1 and len(tokens[0]) <= 4 and tokens[0].endswith(("a", "e", "i", "o", "u")):
        return True

    return False


def _hebrew_word_variant_candidates(raw_word):
    clean_word = _normalize_lookup_word(raw_word)
    letters_only = re.sub(r"[^\u05D0-\u05EA\s]", "", clean_word).strip()
    if not letters_only:
        return []

    variants = []

    def add_variant(value):
        token = str(value or "").strip()
        if token and token not in variants:
            variants.append(token)

    add_variant(letters_only)
    parts = [part for part in letters_only.split() if part]
    if parts:
        add_variant(parts[0])

    for part in parts[:2] if parts else [letters_only]:
        if len(part) >= 4 and part[0] in {"ו", "ה", "ב", "כ", "ל", "מ", "ש"}:
            add_variant(part[1:])

    return variants[:4]


def _parse_meaning_candidates(raw_meaning):
    value = re.sub(r"\s+", " ", str(raw_meaning or "").strip())
    if not value:
        return []

    if not any(sep in value for sep in [";", "/", "|"]):
        return [value]

    chunks = [chunk.strip(" .") for chunk in re.split(
        r"[;/|]", value) if chunk.strip()]
    return chunks[:4] if chunks else [value]


def _collect_word_meaning_alternatives(raw_word, primary_meaning, word_is_hebrew):
    options = []

    def add_option(candidate):
        normalized = re.sub(r"\s+", " ", str(candidate or "").strip(" ."))
        if not normalized:
            return
        if word_is_hebrew and _looks_like_transliteration(normalized):
            return
        lowered = normalized.lower()
        if lowered in {item.lower() for item in options}:
            return
        options.append(normalized)

    if word_is_hebrew:
        variants = _hebrew_word_variant_candidates(raw_word)
        for variant in variants:
            for interpreted in HEBREW_INTERPRETIVE_GLOSSARY.get(variant, []):
                add_option(interpreted)

        for candidate in _parse_meaning_candidates(primary_meaning):
            add_option(candidate)

        for variant in variants[:2]:
            translated, _ = _translate_hebrew_text_online(variant)
            add_option(translated)
    else:
        for candidate in _parse_meaning_candidates(primary_meaning):
            add_option(candidate)

    return options[:3]


def _lookup_hebrew_word_meaning(word):
    def _strip_common_hebrew_prefix(token):
        value = str(token or "").strip()
        if len(value) < 4:
            return value
        if value and value[0] in {"ו", "ה", "ב", "כ", "ל", "מ", "ש"}:
            return value[1:]
        return value

    def _hebrew_word_variants(raw_word):
        normalized = _normalize_lookup_word(raw_word)
        variants = []

        def add_variant(candidate):
            value = str(candidate or "").strip()
            if value and value not in variants:
                variants.append(value)

        add_variant(normalized)
        letters_only = re.sub(r"[^\u05D0-\u05EA\s]", "", normalized).strip()
        add_variant(letters_only)

        if " " in letters_only:
            for part in letters_only.split():
                add_variant(part)
                add_variant(_strip_common_hebrew_prefix(part))
        else:
            add_variant(_strip_common_hebrew_prefix(letters_only))

        return variants[:8]

    clean_word = _normalize_lookup_word(word)
    if not clean_word:
        return "", ""

    variants = _hebrew_word_variants(clean_word)

    for variant in variants:
        exact = HEBREW_WORD_GLOSSARY.get(variant)
        if exact:
            normalized = _normalize_glossary_meaning(exact)
            if normalized:
                return normalized, "local-hebrew-glossary"

    for variant in variants:
        generated, source = _translate_hebrew_text_online(variant)
        if generated and not _looks_like_transliteration(generated):
            return generated, source or "automatic-translation"

    generated, source = _translate_hebrew_text_online(clean_word)
    if generated and not _looks_like_transliteration(generated):
        return generated, source or "automatic-translation"

    return "", ""


def _sanitize_answer_mode(mode_value):
    mode = (mode_value or "balanced").strip().lower()
    return mode if mode in ANSWER_MODES else "balanced"


DETAIL_REQUEST_RE = re.compile(
    r"(\bexplain\b|\bfull\s+explanation\b|\bin\s+depth\b|\bdetailed\b|\bdetail\b|"
    r"\belaborate\b|\bexpand\b|\bbreak\s+down\b|\bwalk\s+me\s+through\b|\bwhy\b|\bhow\b|"
    r"הסבר|למה|כיצד|בפירוט|הרחב|נמק|פרט)",
    re.IGNORECASE,
)


def _is_detail_requested(question, mode):
    mode_value = str(mode or "").strip().lower()
    if mode_value in {"sources", "strict"}:
        return True
    return bool(DETAIL_REQUEST_RE.search(str(question or "")))


def _summarize_ruling_text(ruling_text, max_sentences=3, max_chars=380):
    clean = re.sub(r"\s+", " ", str(ruling_text or "").strip())
    if not clean:
        return ""

    sentence_candidates = [
        segment.strip()
        for segment in re.split(r"(?<=[.!?])\s+", clean)
        if segment.strip()
    ]
    summary = " ".join(sentence_candidates[:max_sentences]).strip(
    ) if sentence_candidates else clean
    if len(summary) > max_chars:
        summary = f"{summary[:max_chars].rstrip()}..."
    return summary


def _extract_action_steps_from_ruling(ruling_text, max_steps=5):
    clean = re.sub(r"\s+", " ", str(ruling_text or "").strip())
    if not clean:
        return []

    numbered_clauses = [
        match.strip(" ;.")
        for match in re.findall(r"\(\d+\)\s*([^;]+)", clean)
        if match and match.strip()
    ]
    if numbered_clauses:
        normalized_steps = []
        for clause in numbered_clauses[:max_steps]:
            clause_clean = clause[0].upper() + clause[1:] if clause else clause
            if clause_clean and clause_clean not in normalized_steps:
                normalized_steps.append(clause_clean)
        if normalized_steps:
            return normalized_steps

    fragments = [
        fragment.strip(" ;")
        for fragment in re.split(r"(?<=[.!?;:])\s+", clean)
        if fragment and fragment.strip()
    ]

    steps = []
    for fragment in fragments:
        lowered = fragment.lower()
        if any(token in lowered for token in (
            "consult",
            "verify",
            "follow",
            "avoid",
            "wait",
            "check",
            "ask",
            "review",
            "custom",
            "practice",
            "מנהג",
            "בדוק",
            "התייעץ",
        )):
            step = fragment[:180].strip()
            if step and step not in steps:
                steps.append(step)
        if len(steps) >= max_steps:
            break

    if steps:
        return steps

    fallback = []
    for fragment in fragments[:max_steps]:
        snippet = fragment[:180].strip()
        if snippet:
            fallback.append(snippet)
    return fallback


def _decode_jsonish_text(value):
    text = str(value or "")
    if not text:
        return ""

    decoded = text.replace("\\n", "\n").replace(
        "\\t", " ").replace("\\\"", '"')
    decoded = re.sub(r"\s+", " ", decoded).strip()
    return decoded


def _extract_jsonish_string_field(text, field_name):
    raw = str(text or "")
    if not raw:
        return ""

    escaped_field = re.escape(str(field_name or "").strip())
    if not escaped_field:
        return ""

    candidates = [raw, raw.replace('\\\"', '"')]
    for candidate in candidates:
        match = re.search(
            rf'"{escaped_field}"\s*:\s*"((?:\\\\.|[^"\\\\])*)"',
            candidate,
            flags=re.IGNORECASE | re.DOTALL,
        )
        if not match:
            continue
        value = _decode_jsonish_text(match.group(1))
        if value:
            return value

    return ""


def _extract_jsonish_string_array_field(text, field_name, max_items=6):
    raw = str(text or "")
    if not raw:
        return []

    escaped_field = re.escape(str(field_name or "").strip())
    if not escaped_field:
        return []

    candidates = [raw, raw.replace('\\\"', '"')]
    for candidate in candidates:
        match = re.search(
            rf'"{escaped_field}"\s*:\s*\[(.*?)\]',
            candidate,
            flags=re.IGNORECASE | re.DOTALL,
        )
        if not match:
            continue

        body = match.group(1)
        extracted = []
        for item in re.findall(r'"((?:\\.|[^"\\])*)"', body, flags=re.DOTALL):
            cleaned = _decode_jsonish_text(item)
            if cleaned and cleaned not in extracted:
                extracted.append(cleaned)
            if len(extracted) >= max_items:
                break

        if extracted:
            return extracted

    return []


def _looks_like_leaked_structured_payload(text):
    normalized = str(text or "").lower()
    if not normalized:
        return False

    leak_markers = (
        "```json",
        '"ruling"',
        '"summary"',
        "## summary",
        "## practical steps",
    )
    return any(marker in normalized for marker in leak_markers)


def _strip_structured_noise(text):
    cleaned = str(text or "").strip()
    if not cleaned:
        return ""

    cleaned = re.sub(r"```(?:json)?", "", cleaned, flags=re.IGNORECASE)
    cleaned = cleaned.replace("```", "")
    cleaned = re.sub(r"^#+\s*(ruling|summary|practical steps?)\s*$", "", cleaned,
                     flags=re.IGNORECASE | re.MULTILINE)
    cleaned = re.sub(r"\*\*(Prohibited|Permitted)\*\*",
                     "", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned).strip()
    return cleaned


def _coerce_ai_answer_shape(result, question, mode, answer_language="en"):
    """Stabilize model output shape so UI always gets readable sections."""
    if not isinstance(result, dict):
        return result

    result_error = str(result.get("error") or "")
    if result_error.startswith("security_blocked_input") or result_error.startswith("security_blocked_domain"):
        return result

    structured = result.get("structured")
    if isinstance(structured, dict):
        clean_structured = dict(structured)
    else:
        clean_structured = None

    raw_answer = str(result.get("answer") or "").strip()

    if clean_structured:
        existing_ruling = str(clean_structured.get("ruling") or "").strip()
        if _looks_like_leaked_structured_payload(existing_ruling) or _looks_like_leaked_structured_payload(raw_answer):
            extracted_ruling = (
                _extract_jsonish_string_field(raw_answer, "ruling")
                or _extract_jsonish_string_field(existing_ruling, "ruling")
            )
            extracted_summary = (
                _extract_jsonish_string_field(raw_answer, "summary")
                or _extract_jsonish_string_field(existing_ruling, "summary")
            )
            extracted_steps = (
                _extract_jsonish_string_array_field(
                    raw_answer, "practical_steps")
                or _extract_jsonish_string_array_field(existing_ruling, "practical_steps")
            )

            if extracted_ruling:
                clean_structured["ruling"] = extracted_ruling
            if extracted_summary:
                clean_structured["summary"] = extracted_summary
            if extracted_steps:
                clean_structured["practical_steps"] = extracted_steps

    if (not clean_structured) or not str(clean_structured.get("ruling") or "").strip():
        parsed = claude.parse_structured_model_output(raw_answer)
        if isinstance(parsed, dict) and str(parsed.get("ruling") or "").strip():
            clean_structured = parsed

    if not clean_structured:
        return result

    detail_needed = _is_detail_requested(question, mode)
    ruling_text = _strip_structured_noise(clean_structured.get("ruling") or "")
    clean_structured["ruling"] = ruling_text

    summary_text = _strip_structured_noise(
        clean_structured.get("summary") or "")
    if summary_text:
        clean_structured["summary"] = summary_text

    practical_steps = clean_structured.get("practical_steps")
    if not isinstance(practical_steps, list):
        practical_steps = []

    if clean_structured.get("is_prohibited") and ruling_text:
        lowered_ruling = ruling_text.lower()
        direct_prohibition = bool(re.search(
            r"(\b(?:not\s+permitted|may\s+not|must\s+not|assur|asur)\b|"
            r"\b(?:is|are|remains|considered|deemed)\s+(?:strictly\s+)?(?:forbidden|prohibited)\b|"
            r"אסור)",
            lowered_ruling,
            flags=re.IGNORECASE,
        ))
        contextual_forbidden_mentions = bool(re.search(
            r"(\bforbidden\s+work\b|\bmelacha\b|\bavoid\s+melacha\b)",
            lowered_ruling,
            flags=re.IGNORECASE,
        ))
        permission_signals = bool(re.search(
            r"(\bpermitted\b|\ballowed\b|\bmitzvah\b|\bobligation\b|\brecommended\b|מותר)",
            lowered_ruling,
            flags=re.IGNORECASE,
        ))

        if not direct_prohibition or (contextual_forbidden_mentions and permission_signals):
            clean_structured["is_prohibited"] = False

    if detail_needed and not summary_text:
        clean_structured["summary"] = _summarize_ruling_text(
            ruling_text,
            max_sentences=4,
            max_chars=520,
        )

    if detail_needed and len(practical_steps) < 2:
        clean_structured["practical_steps"] = _extract_action_steps_from_ruling(
            ruling_text)

    rendered_answer = claude.render_structured_markdown(
        clean_structured,
        answer_language=answer_language,
    )
    if rendered_answer:
        result["structured"] = clean_structured
        result["answer"] = rendered_answer

    return result


def _canonicalize_community_name(name):
    if not name:
        return None

    if name in COMMUNITIES:
        return name

    lowered = name.strip().lower()
    if lowered in COMMUNITY_ALIASES:
        return COMMUNITY_ALIASES[lowered]

    normalized = "".join(ch for ch in lowered if ch.isalnum())
    for alias, canonical in COMMUNITY_ALIASES.items():
        alias_norm = "".join(ch for ch in alias.lower() if ch.isalnum())
        if alias_norm == normalized:
            return canonical

    for canonical in COMMUNITIES.keys():
        canonical_norm = "".join(
            ch for ch in canonical.lower() if ch.isalnum())
        if canonical_norm == normalized:
            return canonical

    return None


def _detect_community_in_text(question):
    q_lower = (question or "").lower()
    for alias, canonical in sorted(COMMUNITY_ALIASES.items(), key=lambda x: len(x[0]), reverse=True):
        if alias in q_lower:
            return canonical

    for canonical in COMMUNITIES.keys():
        if canonical.lower() in q_lower:
            return canonical

    return None


def _build_pyluach_holiday_events(year):
    """Fallback holiday event list for FullCalendar when Hebcal is unavailable."""
    events = []
    try:
        current = greg_date(int(year), 1, 1)
        end = greg_date(int(year), 12, 31)
    except Exception:
        return events

    while current <= end:
        try:
            heb = pyluach_dates.GregorianDate(
                current.year, current.month, current.day).to_heb()
            holiday_name = heb.holiday()
            if holiday_name:
                emoji = _holiday_emoji_for_event(holiday_name, "major")
                events.append({
                    "title": f"{emoji} {holiday_name}",
                    "start": current.isoformat(),
                    "allDay": True,
                    "display": "block",
                    "color": "#802f3e",
                    "textColor": "#ffffff",
                })
        except Exception:
            # Keep fallback generation resilient even if one date fails.
            pass
        current += timedelta(days=1)

    return events


def _strip_leading_symbol_prefix(text):
    raw = str(text or "").strip()
    if not raw:
        return ""
    return re.sub(r"^[^\w\u0590-\u05FF]+", "", raw).strip()


def _holiday_emoji_for_event(title, category=""):
    lowered = str(title or "").strip().lower()
    cat = str(category or "").strip().lower()

    if "yom ha'atzmaut" in lowered or "yom haatzmaut" in lowered:
        return "🇮🇱"
    if "hanukkah" in lowered or "chanukah" in lowered:
        return "🕎"
    if "erev rosh hashana" in lowered or "rosh hashana" in lowered:
        return "🍎🍯"
    if "lag ba'omer" in lowered or "lag baomer" in lowered:
        return "🔥"
    if "yom yerushalayim" in lowered:
        return "🇮🇱"
    if "erev shavuot" in lowered:
        return "⛰️"
    if "shavuot" in lowered:
        return "🌸"
    if any(token in lowered for token in ("sukkot", "succot", "sukkos", "succos")):
        return "🍋🌿"
    if "rosh chodesh" in lowered or cat == "roshchodesh":
        return "🌙"
    if cat == "fast" or any(token in lowered for token in (
        "taanis", "taanit", "fast", "tzom", "tisha b'av", "17 of tamuz", "gedaliah", "esther"
    )):
        return "✡️"
    if "shabbat" in lowered or cat in {"shabbat", "parashat"}:
        return "🕍"
    if cat in {"modern"}:
        return "✡️"
    if cat in {"major", "minor", "holiday", "special"}:
        return "✡️"
    return "✡️"


def _holiday_color_for_category(category):
    palette = {
        "major": "#802f3e",
        "minor": "#594176",
        "modern": "#2563eb",
        "fast": "#374151",
        "roshchodesh": "#5a99b7",
        "shabbat": "#004e5f",
        "parashat": "#004e5f",
        "holiday": "#802f3e",
        "special": "#6b7280",
    }
    return palette.get(str(category or "").strip().lower(), "#6b7280")


load_dotenv()
app = Flask(__name__)

# Configure structured JSON logging as early as possible so all log records
# (including import-time warnings from sub-modules) use the JSON formatter.
setup_logging()

_flask_secret = os.environ.get("FLASK_SECRET_KEY")
if not _flask_secret:
    import logging as _logging
    _logging.getLogger(__name__).critical(
        "FLASK_SECRET_KEY is not set — using a random ephemeral key. "
        "All sessions will be invalidated on every process restart. "
        "Set FLASK_SECRET_KEY in your environment for a stable key."
    )
    _flask_secret = os.urandom(32)
app.secret_key = _flask_secret

RATE_LIMIT_DEFAULT = [
    item.strip()
    for item in (os.environ.get("RATE_LIMIT_DEFAULT") or "").split(",")
    if item.strip()
]
RATE_LIMIT_ASK = (os.environ.get("RATE_LIMIT_ASK") or "20 per minute").strip()
RATE_LIMIT_STORAGE_URI = (os.environ.get(
    "RATELIMIT_STORAGE_URI") or "memory://").strip()


def _rate_limit_key():
    forwarded = (request.headers.get("CF-Connecting-IP")
                 or request.headers.get("X-Forwarded-For") or "").split(",")[0].strip()
    real_ip = (request.headers.get("X-Real-IP") or "").strip()
    return forwarded or real_ip or request.remote_addr or "127.0.0.1"


limiter = None
if Limiter is not None:
    limiter_kwargs = {
        "app": app,
        "key_func": _rate_limit_key,
        "storage_uri": RATE_LIMIT_STORAGE_URI,
    }
    if RATE_LIMIT_DEFAULT:
        limiter_kwargs["default_limits"] = RATE_LIMIT_DEFAULT
    limiter = Limiter(**limiter_kwargs)


def maybe_limit(limit_value):
    """Apply rate limiting when Flask-Limiter is available."""
    def decorator(route_fn):
        if limiter is None or not limit_value:
            return route_fn
        return limiter.limit(limit_value)(route_fn)

    return decorator


CLERK_PUBLISHABLE_KEY = (
    os.environ.get("CLERK_PUBLISHABLE_KEY")
    or os.environ.get("NEXT_PUBLIC_CLERK_PUBLISHABLE_KEY")
    or ""
).strip()
CLERK_JWT_ISSUER = (os.environ.get("CLERK_JWT_ISSUER")
                    or "").strip().rstrip("/")
CLERK_AUDIENCE = (os.environ.get("CLERK_AUDIENCE") or "").strip()
_in_prod_runtime = (
    os.environ.get("VERCEL") == "1"
    or os.environ.get("FLASK_ENV", "").strip().lower() == "production"
)
# Default: enforce auth on Vercel/production, allow unauthenticated on local dev.
CLERK_ENFORCE_AUTH = (
    os.environ.get("CLERK_ENFORCE_AUTH")
    or ("true" if _in_prod_runtime else "false")
).strip().lower() == "true"
_clerk_jwks_client = None

SUPABASE_URL = (os.environ.get("SUPABASE_URL") or "").strip()
if not SUPABASE_URL:
    SUPABASE_URL = (os.environ.get("NEXT_PUBLIC_SUPABASE_URL") or "").strip()

SUPABASE_PUBLISHABLE_KEY = (
    os.environ.get("SUPABASE_ANON_KEY")
    or os.environ.get("NEXT_PUBLIC_SUPABASE_PUBLISHABLE_DEFAULT_KEY")
    or os.environ.get("NEXT_PUBLIC_SUPABASE_ANON_KEY")
    or ""
).strip()

SUPABASE_SERVICE_ROLE_KEY = (os.environ.get(
    "SUPABASE_SERVICE_ROLE_KEY") or "").strip()
SUPABASE_PREFS_TABLE = (os.environ.get(
    "SUPABASE_PREFS_TABLE") or "user_preferences").strip()
SUPABASE_COMMUNITY_KNOWLEDGE_TABLE = (os.environ.get(
    "SUPABASE_COMMUNITY_KNOWLEDGE_TABLE") or "community_knowledge").strip()
SUPABASE_USER_MEMORIES_TABLE = (os.environ.get(
    "SUPABASE_USER_MEMORIES_TABLE") or "user_memories").strip()
SUPABASE_STUDY_BOOKMARKS_TABLE = (os.environ.get(
    "SUPABASE_STUDY_BOOKMARKS_TABLE") or "study_bookmarks").strip()
STRICT_SUPABASE_RLS = (os.environ.get("STRICT_SUPABASE_RLS")
                       or "true").strip().lower() == "true"
ERROR_LOG_WEBHOOK_URL = (os.environ.get("ERROR_LOG_WEBHOOK_URL") or "").strip()
_supabase_client = None


def _env_int(name, default):
    try:
        raw_value = os.environ.get(name)
        if raw_value is None:
            return int(default)
        return int(raw_value)
    except Exception:
        return default


RAG_TOP_KNOWLEDGE_ROWS = _env_int("RAG_TOP_KNOWLEDGE_ROWS", 5)
RAG_MEMORY_ROWS = _env_int("RAG_MEMORY_ROWS", 2)


def _extract_bearer_token():
    auth_header = request.headers.get("Authorization", "")
    if not auth_header.lower().startswith("bearer "):
        return None
    token = auth_header.split(" ", 1)[1].strip()
    return token or None


def _get_clerk_jwks_client():
    global _clerk_jwks_client
    if not CLERK_JWT_ISSUER:
        return None
    if _clerk_jwks_client is None:
        jwks_url = f"{CLERK_JWT_ISSUER}/.well-known/jwks.json"
        _clerk_jwks_client = jwt.PyJWKClient(jwks_url)
    return _clerk_jwks_client


def _verify_clerk_token(token):
    if not token:
        raise ValueError("Missing bearer token")
    if not CLERK_JWT_ISSUER:
        raise ValueError("Server missing CLERK_JWT_ISSUER")

    jwks_client = _get_clerk_jwks_client()
    if jwks_client is None:
        raise ValueError("Clerk JWKS client unavailable")

    signing_key = jwks_client.get_signing_key_from_jwt(token).key
    decode_kwargs = {
        "algorithms": ["RS256"],
        "issuer": CLERK_JWT_ISSUER,
    }
    if CLERK_AUDIENCE:
        decode_kwargs["audience"] = CLERK_AUDIENCE
    else:
        decode_kwargs["options"] = {"verify_aud": False}

    return jwt.decode(token, signing_key, **decode_kwargs)


def _get_supabase_client():
    global _supabase_client
    if create_client is None:
        return None
    if not SUPABASE_URL or not SUPABASE_SERVICE_ROLE_KEY:
        return None
    if _supabase_client is None:
        # Backend retrieval uses service role to avoid RLS limits on anon/auth keys.
        _supabase_client = create_client(
            SUPABASE_URL, SUPABASE_SERVICE_ROLE_KEY)
    return _supabase_client


def _looks_like_jwt(value):
    if not isinstance(value, str):
        return False
    parts = value.split(".")
    return len(parts) == 3 and all(parts)


def _extract_supabase_token_from_cookie_value(raw_value):
    if not raw_value:
        return None

    decoded = unquote(raw_value)
    if _looks_like_jwt(decoded):
        return decoded

    try:
        parsed = json.loads(decoded)
    except Exception:
        return None

    if isinstance(parsed, dict):
        token = parsed.get("access_token") or parsed.get("accessToken")
        return token if isinstance(token, str) and token else None

    if isinstance(parsed, list):
        for item in parsed:
            if isinstance(item, str) and _looks_like_jwt(item):
                return item
            if isinstance(item, dict):
                token = item.get("access_token") or item.get("accessToken")
                if isinstance(token, str) and token:
                    return token

    return None


def _extract_supabase_access_token():
    # Prefer Authorization header so API clients can override cookie auth.
    bearer = _extract_bearer_token()
    if bearer:
        return bearer

    direct_cookie_names = [
        "sb-access-token",
        "supabase-access-token",
    ]
    for cookie_name in direct_cookie_names:
        direct_value = request.cookies.get(cookie_name)
        token = _extract_supabase_token_from_cookie_value(direct_value)
        if token:
            return token

    session_cookie_values = []
    chunked_cookies = {}
    for cookie_name, cookie_value in request.cookies.items():
        if not (cookie_name.startswith("sb-") and "-auth-token" in cookie_name):
            continue

        if "." in cookie_name:
            base, suffix = cookie_name.rsplit(".", 1)
            if suffix.isdigit():
                chunked_cookies.setdefault(base, []).append(
                    (int(suffix), cookie_value))
                continue

        session_cookie_values.append(cookie_value)

    for cookie_value in session_cookie_values:
        token = _extract_supabase_token_from_cookie_value(cookie_value)
        if token:
            return token

    for _, chunk_parts in chunked_cookies.items():
        sorted_parts = sorted(chunk_parts, key=lambda part: part[0])
        joined_value = "".join(part[1] for part in sorted_parts)
        token = _extract_supabase_token_from_cookie_value(joined_value)
        if token:
            return token

    return None


def _get_request_supabase_client():
    """Flask equivalent of Next.js createServerClient for request-scoped reads."""
    if create_client is None:
        return None
    if not SUPABASE_URL or not SUPABASE_PUBLISHABLE_KEY:
        return None

    access_token = _extract_supabase_access_token()
    if not access_token or SyncClientOptions is None:
        return create_client(SUPABASE_URL, SUPABASE_PUBLISHABLE_KEY)

    auth_headers = {"Authorization": f"Bearer {access_token}"}
    try:
        options = SyncClientOptions(headers=auth_headers)
        return create_client(SUPABASE_URL, SUPABASE_PUBLISHABLE_KEY, options=options)
    except TypeError:
        # Compatibility fallback for older supabase-py signatures.
        return create_client(SUPABASE_URL, SUPABASE_PUBLISHABLE_KEY)


def _get_user_scoped_supabase_client():
    """Return request-scoped Supabase client for RLS-protected user tables."""
    client = _get_request_supabase_client()
    if not client:
        return None

    if STRICT_SUPABASE_RLS and not _extract_supabase_access_token():
        return None

    return client


def _capture_backend_error(event_name, error, context=None):
    """Sentry-style structured logger for backend failures and AI prompt issues."""
    context = context if isinstance(context, dict) else {}
    message = str(error) if error is not None else ""
    payload = {
        "event": str(event_name or "unknown"),
        "message": message,
        "context": context,
        "ts": int(time.time()),
    }

    app.logger.error("OBS_EVENT %s", json.dumps(payload, ensure_ascii=True),
                     exc_info=error if isinstance(error, Exception) else False)

    if ERROR_LOG_WEBHOOK_URL:
        try:
            requests.post(
                ERROR_LOG_WEBHOOK_URL,
                json=payload,
                timeout=2,
            )
        except Exception:
            pass


def _get_request_user_id():
    claims = getattr(g, "clerk_claims", {}) or {}
    user_id = str(claims.get("sub") or "").strip()
    if user_id:
        return user_id

    token = _extract_bearer_token()
    if not token:
        return None

    try:
        decoded = _verify_clerk_token(token)
    except Exception:
        return None

    user_id = str(decoded.get("sub") or "").strip()
    return user_id or None


def _normalize_rag_text(value, max_chars=360):
    text = re.sub(r"\s+", " ", str(value or "").strip())
    if max_chars > 0 and len(text) > max_chars:
        text = f"{text[:max_chars].rstrip()}..."
    return text


def _score_community_knowledge_row(row, keywords, canonical_lens):
    topic = str(row.get("topic") or "").lower()
    source = str(row.get("halakhic_source") or "").lower()
    content = str(row.get("content") or "").lower()
    community_name = str(row.get("community_name") or "").lower()

    score = 0
    if canonical_lens and canonical_lens != "All":
        lens_text = canonical_lens.lower().strip()
        if lens_text and lens_text in community_name:
            score += 8

    for keyword in keywords:
        if keyword in topic:
            score += 8
        if keyword in source:
            score += 4
        if keyword in content:
            score += 2

    return score


def _community_filter_from_request(query, canonical_lens):
    if canonical_lens and canonical_lens != "All":
        return canonical_lens

    detected = _detect_community_in_text(query)
    return detected or None


def _build_knowledge_text_or_filter(keywords, max_keywords=6):
    conditions = []
    for keyword in (keywords or [])[:max_keywords]:
        clean = re.sub(r"[^A-Za-z0-9\u0590-\u05FF\-]",
                       "", str(keyword or "").strip())
        if not clean:
            continue

        pattern = f"%{clean}%"
        conditions.extend([
            f"topic.ilike.{pattern}",
            f"content.ilike.{pattern}",
        ])

    return ",".join(conditions)


def _retrieve_community_knowledge(query, canonical_lens="All", max_rows=None):
    supabase = _get_supabase_client()
    if not supabase:
        return []

    target_rows = max_rows or RAG_TOP_KNOWLEDGE_ROWS
    keywords = _extract_query_keywords(query, max_keywords=10)
    community_filter = _community_filter_from_request(query, canonical_lens)
    text_or_filter = _build_knowledge_text_or_filter(keywords)

    query_row_cap = max(50, min(600, target_rows * 25))

    try:
        def run_query(apply_text_filter=True):
            table = supabase.table(SUPABASE_COMMUNITY_KNOWLEDGE_TABLE).select(
                "id,community_name,topic,halakhic_source,content"
            )

            if community_filter:
                # Case-insensitive match so Ashkenaz/Ashkenazi variants still return rows.
                table = table.ilike("community_name", f"%{community_filter}%")

            if apply_text_filter and text_or_filter:
                table = table.or_(text_or_filter)

            result = table.limit(query_row_cap).execute()
            return result.data if isinstance(result.data, list) else []

        rows = run_query(apply_text_filter=True)

        # Fallback to community-only retrieval when text filter is too restrictive.
        if not rows and community_filter and text_or_filter:
            rows = run_query(apply_text_filter=False)
    except Exception:
        return []

    ranked = []
    for row in rows:
        if not isinstance(row, dict):
            continue

        score = _score_community_knowledge_row(
            row,
            keywords,
            community_filter or canonical_lens,
        )
        if score <= 0 and keywords and text_or_filter:
            continue

        ranked.append((score, row))

    ranked.sort(
        key=lambda item: (
            item[0],
            len(str(item[1].get("topic") or "")),
        ),
        reverse=True,
    )

    top_rows = []
    for score, row in ranked[:target_rows]:
        top_rows.append({
            "id": str(row.get("id") or "").strip(),
            "community_name": str(row.get("community_name") or "").strip(),
            "topic": str(row.get("topic") or "").strip(),
            "halakhic_source": str(row.get("halakhic_source") or "").strip(),
            "content": _normalize_rag_text(row.get("content")),
            "score": score,
        })

    return top_rows


def _knowledge_rows_to_customs(rows):
    customs = []
    for row in rows:
        if not isinstance(row, dict):
            continue

        topic = str(row.get("topic") or "").strip()
        source = str(row.get("halakhic_source") or "").strip()
        content = str(row.get("content") or "").strip()
        customs.append({
            "community": str(row.get("community_name") or "").strip(),
            "topic": topic,
            "ruling": content,
            "source": source,
            "notes": f"Topic: {topic}" if topic else "",
        })

    return customs


def _fetch_user_memory_summaries(user_id, limit=None):
    if not user_id:
        return []

    supabase = _get_user_scoped_supabase_client()
    if not supabase and not STRICT_SUPABASE_RLS:
        supabase = _get_supabase_client()
    if not supabase:
        return []

    target_limit = limit or RAG_MEMORY_ROWS
    try:
        result = (
            supabase
            .table(SUPABASE_USER_MEMORIES_TABLE)
            .select("summary,created_at")
            .eq("user_id", user_id)
            .order("created_at", desc=True)
            .limit(target_limit)
            .execute()
        )
    except Exception:
        return []

    rows = result.data if isinstance(result.data, list) else []
    summaries = []
    for row in rows:
        if not isinstance(row, dict):
            continue

        summary = _normalize_rag_text(row.get("summary"), max_chars=260)
        if not summary:
            continue

        summaries.append({
            "summary": summary,
            "created_at": row.get("created_at"),
        })

    return summaries


def _build_interaction_summary(question, answer):
    clean_q = _normalize_rag_text(question, max_chars=160)
    clean_a = re.sub(r"[#*_`~>\-]+", " ", str(answer or ""))
    clean_a = _normalize_rag_text(clean_a, max_chars=240)
    return f"Q: {clean_q} | A: {clean_a}".strip()


def _store_user_memory_summary(user_id, question, answer):
    if not user_id:
        return

    supabase = _get_user_scoped_supabase_client()
    if not supabase and not STRICT_SUPABASE_RLS:
        supabase = _get_supabase_client()
    if not supabase:
        return

    summary = _build_interaction_summary(question, answer)
    if not summary:
        return

    payload = {
        "id": str(uuid4()),
        "user_id": user_id,
        "summary": summary,
    }

    try:
        supabase.table(SUPABASE_USER_MEMORIES_TABLE).insert(payload).execute()
    except Exception:
        # Memory write failures should never block the user response path.
        return


def maybe_require_clerk_auth(route_fn):
    @wraps(route_fn)
    def wrapped(*args, **kwargs):
        token = _extract_bearer_token()
        if not token:
            if CLERK_ENFORCE_AUTH:
                return jsonify({"error": "Authentication required"}), 401
            return route_fn(*args, **kwargs)

        try:
            g.clerk_claims = _verify_clerk_token(token)
        except Exception:
            return jsonify({"error": "Invalid or expired Clerk token"}), 401

        return route_fn(*args, **kwargs)

    return wrapped


def require_clerk_auth(route_fn):
    @wraps(route_fn)
    def wrapped(*args, **kwargs):
        token = _extract_bearer_token()
        if not token:
            return jsonify({"error": "Authentication required"}), 401

        try:
            g.clerk_claims = _verify_clerk_token(token)
        except Exception:
            return jsonify({"error": "Invalid or expired Clerk token"}), 401

        return route_fn(*args, **kwargs)

    return wrapped


def _get_prayer_refs(prayer_name):
    """Resolve prayer/service name to a list of Sefaria refs."""
    resolved_name = (unquote(prayer_name or "") or "").strip()
    if resolved_name in SIDDUR_SECTION_MAP:
        return SIDDUR_SECTION_MAP[resolved_name]

    from backend.sefaria_library import get_index_leaf_refs
    return get_index_leaf_refs(resolved_name, max_refs=80)


def _coerce_coordinate(value, min_value, max_value):
    try:
        numeric = float(value)
    except (TypeError, ValueError):
        return None
    if numeric < min_value or numeric > max_value:
        return None
    return numeric


def _coerce_int(value, default, min_value=1, max_value=100):
    try:
        parsed = int(value)
    except (TypeError, ValueError):
        return default
    return max(min_value, min(parsed, max_value))


# Browser cache/session cadence defaults to 2 days so clients revalidate often.
RESOURCE_RELOAD_SECONDS = _coerce_int(
    os.environ.get("RESOURCE_RELOAD_SECONDS"),
    default=60 * 60 * 24 * 2,
    min_value=60 * 60,
    max_value=60 * 60 * 24 * 14,
)
SESSION_RELOAD_SECONDS = _coerce_int(
    os.environ.get("SESSION_RELOAD_SECONDS"),
    default=RESOURCE_RELOAD_SECONDS,
    min_value=60 * 60,
    max_value=60 * 60 * 24 * 30,
)
STATIC_STALE_WHILE_REVALIDATE_SECONDS = max(
    60 * 60,
    min(60 * 60 * 24, RESOURCE_RELOAD_SECONDS // 2),
)

is_production_runtime = (
    os.environ.get("VERCEL") == "1"
    or os.environ.get("FLASK_ENV", "").strip().lower() == "production"
)

app.config["PERMANENT_SESSION_LIFETIME"] = timedelta(
    seconds=SESSION_RELOAD_SECONDS)
app.config["SESSION_REFRESH_EACH_REQUEST"] = True
app.config["SESSION_COOKIE_HTTPONLY"] = True
app.config["SESSION_COOKIE_SAMESITE"] = "Lax"
app.config["SESSION_COOKIE_SECURE"] = is_production_runtime
app.config["SEND_FILE_MAX_AGE_DEFAULT"] = RESOURCE_RELOAD_SECONDS


@app.before_request
def apply_session_cookie_policy():
    # Ensure Flask issues an expiring cookie instead of a browser-session cookie.
    session.permanent = True


@app.after_request
def apply_response_cache_policy(response):
    path = request.path or ""

    # Security headers for every response
    response.headers.setdefault("X-Frame-Options", "SAMEORIGIN")
    response.headers.setdefault("X-Content-Type-Options", "nosniff")
    response.headers.setdefault(
        "Referrer-Policy", "strict-origin-when-cross-origin")
    response.headers.setdefault("X-XSS-Protection", "1; mode=block")
    response.headers.setdefault(
        "Content-Security-Policy",
        "default-src 'self'; "
        "script-src 'self' 'unsafe-inline' 'unsafe-eval' "
        "https://cdn.tailwindcss.com https://cdn.jsdelivr.net "
        "https://js.clerk.com https://clerk.com; "
        "style-src 'self' 'unsafe-inline' "
        "https://cdn.jsdelivr.net https://fonts.googleapis.com; "
        "font-src 'self' https://fonts.gstatic.com; "
        "img-src 'self' data: https:; "
        "connect-src 'self' https://clerk.com https://*.clerk.accounts.dev "
        "https://api.clerk.com; "
        "frame-ancestors 'none'; "
        "base-uri 'self'; "
        "form-action 'self';"
    )

    if path.startswith("/api/") or path in {"/ask", "/set_location"}:
        response.headers["Cache-Control"] = "no-store"
        return response

    if path == "/service-worker.js":
        response.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
        response.headers["Pragma"] = "no-cache"
        response.headers["Expires"] = "0"
        return response

    if path.startswith("/static/"):
        response.headers["Cache-Control"] = (
            f"public, max-age={RESOURCE_RELOAD_SECONDS}, "
            f"stale-while-revalidate={STATIC_STALE_WHILE_REVALIDATE_SECONDS}"
        )
        return response

    if path == "/manifest.webmanifest":
        response.headers["Cache-Control"] = (
            f"public, max-age={RESOURCE_RELOAD_SECONDS}, must-revalidate"
        )
        return response

    if response.mimetype in {"text/html", "application/xhtml+xml"}:
        response.headers["Cache-Control"] = (
            f"public, max-age={RESOURCE_RELOAD_SECONDS}, must-revalidate"
        )

    return response


def _parse_multi_value_arg(name):
    raw = (request.args.get(name, "") or "").strip()
    if not raw:
        return []
    parts = []
    for chunk in raw.split(","):
        value = chunk.strip()
        if value:
            parts.append(value)
    return parts


def _extract_search_metadata_filters():
    metadata_filters = {}
    for key in ("era", "author", "category", "geography", "nusach"):
        values = _parse_multi_value_arg(key)
        if values:
            metadata_filters[key] = values
    return metadata_filters


def _extract_client_ip():
    forwarded_for = (request.headers.get("X-Forwarded-For")
                     or "").split(",")[0].strip()
    if forwarded_for:
        return forwarded_for

    real_ip = (request.headers.get("X-Real-IP") or "").strip()
    if real_ip:
        return real_ip

    remote_ip = (request.remote_addr or "").strip()
    return remote_ip or None


def get_engine():
    # Instantiate engine using session location or IP fallback
    lat = _coerce_coordinate(session.get('lat'), -90, 90)
    lon = _coerce_coordinate(session.get('lon'), -180, 180)

    if lat is None or lon is None:
        client_ip = _extract_client_ip()
        ip_target = ""
        if client_ip and client_ip not in {"127.0.0.1", "::1"}:
            ip_target = client_ip

        try:
            # ip-api.com is free, no key required, ~45 req/min limit.
            # Use request IP from Vercel headers instead of server runtime IP.
            # ip-api free tier only supports HTTP; ipwho.is is HTTPS fallback.
            lookup_urls = [
                f"http://ip-api.com/json/{ip_target}?fields=status,lat,lon,timezone,query",
                f"https://ipwho.is/{ip_target}" if ip_target else "https://ipwho.is/",
            ]

            for lookup_url in lookup_urls:
                r = requests.get(lookup_url, timeout=3)
                data = r.json() if r.ok else {}

                ip_lat = None
                ip_lon = None
                if data.get("status") == "success":
                    ip_lat = _coerce_coordinate(data.get('lat'), -90, 90)
                    ip_lon = _coerce_coordinate(data.get('lon'), -180, 180)
                elif data.get("success") is True:
                    ip_lat = _coerce_coordinate(data.get('latitude'), -90, 90)
                    ip_lon = _coerce_coordinate(
                        data.get('longitude'), -180, 180)

                if ip_lat is not None and ip_lon is not None:
                    lat = ip_lat
                    lon = ip_lon
                    session['lat'] = lat
                    session['lon'] = lon
                    break
        except Exception as e:
            app.logger.warning(f"Location IP lookup failed: {str(e)}")

    if lat is None or lon is None:
        # Defaulting to NYC if all lookup methods fail
        app.logger.info("Using default location: New York City")
        lat, lon = (40.7128, -74.0060)

    return ShelahEngine(lat=lat, lon=lon)


@app.route("/")
@app.route("/settings")
@app.route("/profile")
def index():
    engine = get_engine()
    daily_study = engine.get_daily_learning()

    # We no longer need hebcal learning as per new architecture, relying on Sefaria cal
    return render_template(
        "index.html",
        daily=daily_study,
        clerk_publishable_key=CLERK_PUBLISHABLE_KEY,
        clerk_enforce_auth=CLERK_ENFORCE_AUTH,
    )


@app.route("/terms")
def terms():
    return render_template(
        "terms.html",
        clerk_publishable_key=CLERK_PUBLISHABLE_KEY,
        clerk_enforce_auth=CLERK_ENFORCE_AUTH,
    )


@app.route("/privacy")
def privacy():
    return render_template(
        "privacy.html",
        clerk_publishable_key=CLERK_PUBLISHABLE_KEY,
        clerk_enforce_auth=CLERK_ENFORCE_AUTH,
    )


@app.route("/api/accept-legal", methods=["POST"])
@maybe_require_clerk_auth
def accept_legal():
    """Record that a user has accepted the Terms of Service and Privacy Policy."""
    user_id = _get_request_user_id()
    if not user_id:
        # Not authenticated — store acceptance in localStorage only (handled client-side).
        return jsonify({"success": True, "stored": "client"}), 200

    try:
        supabase_client = _get_supabase_client()
        if supabase_client:
            supabase_client.table("user_preferences").upsert(
                {
                    "clerk_id": user_id,
                    "legal_accepted": True,
                    "legal_accepted_at": datetime.utcnow().isoformat(),
                },
                on_conflict="clerk_id",
            ).execute()
    except Exception as e:
        app.logger.warning(
            "accept_legal: could not persist to Supabase: %s", e)

    return jsonify({"success": True, "stored": "server"}), 200


@app.route('/set_location', methods=['POST'])
def set_location():
    data = request.get_json(silent=True) or {}
    lat = _coerce_coordinate(data.get('lat'), -90, 90)
    lon = _coerce_coordinate(data.get('lon'), -180, 180)
    if lat is None or lon is None:
        return jsonify({"error": "Invalid coordinates provided. Values must be numeric and within valid lat/lon ranges."}), 400

    session['lat'] = lat
    session['lon'] = lon
    return jsonify({"status": "success", "lat": lat, "lon": lon})


@app.route('/api/zmanim')
def get_zmanim_api():
    community = request.args.get('community', 'standard')
    lat = _coerce_coordinate(request.args.get('lat'), -90, 90)
    lon = _coerce_coordinate(request.args.get('lon'), -180, 180)

    if lat is not None and lon is not None:
        session['lat'] = lat
        session['lon'] = lon
        engine = ShelahEngine(lat=lat, lon=lon)
    else:
        engine = get_engine()

    times = engine.get_zmanim(community)
    return jsonify(times)


@app.route('/api/zmanim/month')
def get_zmanim_month():
    lat = _coerce_coordinate(request.args.get('lat'), -90, 90)
    lon = _coerce_coordinate(request.args.get('lon'), -180, 180)

    if lat is not None and lon is not None:
        session['lat'] = lat
        session['lon'] = lon
        engine = ShelahEngine(lat=lat, lon=lon)
    else:
        engine = get_engine()

    events = engine.get_monthly_zmanim()
    return jsonify(events)


@app.route("/manifest.webmanifest")
def web_manifest():
    return send_from_directory("static", "manifest.webmanifest", mimetype="application/manifest+json")


@app.route("/favicon.ico")
def favicon():
    return send_from_directory("static", "favicon.svg", mimetype="image/svg+xml")


@app.route("/service-worker.js")
def service_worker():
    response = send_from_directory(
        "static", "service-worker.js", mimetype="application/javascript")
    response.headers["Cache-Control"] = "no-cache"
    return response


@app.errorhandler(429)
def handle_rate_limit(_error):
    return jsonify({
        "error": "Rate limit exceeded. Please wait and retry.",
        "status": "rate_limited",
    }), 429


@app.route("/ask", methods=["POST"])
@maybe_require_clerk_auth
@maybe_limit(RATE_LIMIT_ASK)
def ask_question():
    data = request.get_json(silent=True) or {}
    raw_question = data.get("question", "")
    question = claude.sanitize_user_query(raw_question)
    question_was_sanitized = question != str(raw_question or "").strip()

    if not question:
        return jsonify({"error": "No valid question provided"}), 400
    answer_language = str(data.get("language") or "en").strip().lower()
    if answer_language not in {"en", "he"}:
        answer_language = "en"
    mode = _sanitize_answer_mode(data.get("mode"))
    community_lens = (data.get("community") or "All").strip()
    canonical_lens = "All" if community_lens.lower() == "all" else (
        _canonicalize_community_name(community_lens) or community_lens)
    user_id = _get_request_user_id()
    ask_cache_key = "|".join([
        question.lower(),
        answer_language,
        mode,
        canonical_lens.lower(),
        user_id or "anon",
    ])

    cached_payload = _get_cached_ask_payload(ask_cache_key)
    if cached_payload is not None:
        cached_meta = cached_payload.get("meta")
        if isinstance(cached_meta, dict):
            cached_meta["cached"] = True
            cached_meta["generated_at"] = int(time.time())
        return jsonify(cached_payload)

    try:
        engine = get_engine()

        # Check for direct prayer-service questions
        if any(prayer in question for prayer in ["Shacharit", "Mincha", "Maariv", "Kiddush", "Havdalah"]):
            # Return a prayer service focused response
            DEVTOOLS_STATS["answers_total"] += 1
            prayer_answer = (
                f"Prayer Service Guide\n\n{question}\n\n"
                "You can browse full liturgy books and services from the prayer sections. "
                "For practical application, compare local community custom with your rabbi's guidance."
            )
            if answer_language == "he":
                prayer_answer = (
                    f"מדריך תפילה\n\n{question}\n\n"
                    "ניתן לעיין בספרי התפילה והשירותים הליטורגיים המלאים באזור התפילה. "
                    "להכרעה מעשית יש להשוות למנהג הקהילה המקומית ולהתייעץ עם הרב שלך."
                )
            prayer_payload = {
                "answer": prayer_answer,
                "confidence": 0.85,
                "sources": [{
                    "ref": "Sefaria Liturgy",
                    "title": "Sefaria Prayer Books",
                    "lines": [{"en": f"Prayer Service: {question}", "he": f"תפילה: {question}"}]
                }],
                "customs": [],
                "meta": {
                    "mode": mode,
                    "language": answer_language,
                    "community_lens": canonical_lens,
                    "source_count": 1,
                    "custom_count": 0,
                    "generated_at": int(time.time()),
                    "fallback": False,
                    "cached": False,
                }
            }
            _set_cached_ask_payload(ask_cache_key, prayer_payload)
            return jsonify(prayer_payload)

        # 1. Fetch Sefaria Refs - Standard halakhic questions
        primary_refs = sefaria.find_refs_for_question(question)
        max_primary_refs = _env_int(
            "ASK_PRIMARY_SOURCE_LIMIT", 4)  # Capped at 4 for speed
        max_primary_refs = max(1, min(max_primary_refs, 8))
        primary_ref_candidates = []
        for ref in primary_refs:
            normalized_ref = str(ref or "").strip()
            if not normalized_ref:
                continue
            primary_ref_candidates.append(normalized_ref)
            if len(primary_ref_candidates) >= max_primary_refs:
                break

        primary_sources = []
        if primary_ref_candidates:
            worker_count = min(4, len(primary_ref_candidates))
            with ThreadPoolExecutor(max_workers=worker_count) as executor:
                source_futures = [
                    executor.submit(engine.get_library_text, ref)
                    for ref in primary_ref_candidates
                ]
                for future in source_futures:
                    try:
                        source_data = future.result(timeout=3)
                    except Exception:
                        continue
                    if isinstance(source_data, dict):
                        primary_sources.append(source_data)

        # 2-4. Fetch remaining context in parallel.
        with ThreadPoolExecutor(max_workers=4) as executor:
            halachipedia_future = executor.submit(
                engine.get_halachipedia_summary, question)
            knowledge_future = executor.submit(
                _retrieve_community_knowledge,
                question,
                canonical_lens=canonical_lens,
                max_rows=RAG_TOP_KNOWLEDGE_ROWS,
            )
            memory_future = executor.submit(
                _fetch_user_memory_summaries,
                user_id,
                limit=RAG_MEMORY_ROWS,
            )
            wiki_future = executor.submit(engine.get_wiki, question)

            try:
                halachipedia_info = halachipedia_future.result(timeout=4)
            except Exception:
                halachipedia_info = None
            try:
                knowledge_rows = knowledge_future.result(timeout=4)
            except Exception:
                knowledge_rows = []
            try:
                user_memory_summaries = memory_future.result(timeout=5)
            except Exception:
                user_memory_summaries = []
            try:
                wiki_info = wiki_future.result(timeout=3)
            except Exception:
                wiki_info = None

        halachipedia_list = [halachipedia_info] if halachipedia_info else []
        knowledge_rows = knowledge_rows if isinstance(
            knowledge_rows, list) else []
        user_memory_summaries = user_memory_summaries if isinstance(
            user_memory_summaries, list) else []
        customs_info = _knowledge_rows_to_customs(knowledge_rows)
        wiki_list = [wiki_info] if wiki_info else []

        # 5. Prepare flattened primary source text for the protected AI wrapper.
        flat_sources_for_claude = []
        for src in primary_sources:
            src_lines = src.get('lines', []) if isinstance(src, dict) else []
            if not isinstance(src_lines, list):
                src_lines = []
            preferred_lines = []
            for line in src_lines:
                if not isinstance(line, dict):
                    continue
                preferred = (line.get('he') or line.get('en')) if answer_language == 'he' else (
                    line.get('en') or line.get('he'))
                if preferred:
                    preferred_lines.append(str(preferred).strip())
            flat_sources_for_claude.append({
                'ref': str(src.get('ref') or '') if isinstance(src, dict) else '',
                'text': ' '.join(preferred_lines)
            })

        has_primary_sources = bool(flat_sources_for_claude)
        has_customs = bool(knowledge_rows)
        has_whitelisted_external = bool(halachipedia_list)
        use_tertiary_web_context = (
            not has_primary_sources
            and not has_customs
            and not has_whitelisted_external
        )
        wiki_context_for_claude = wiki_list if use_tertiary_web_context else []

        if mode == "strict" and not flat_sources_for_claude:
            DEVTOOLS_STATS["answers_total"] += 1
            DEVTOOLS_STATS["strict_blocks"] += 1
            DEVTOOLS_STATS["fallback_answers"] += 1
            display_sources = _compact_ai_sources(primary_sources)
            strict_payload = {
                "answer": (
                    "Strict Sources Mode could not complete this request because no primary Sefaria sources "
                    "were matched with sufficient confidence. Please refine the question with a text reference."
                ),
                "confidence": 0.2,
                "wiki": wiki_list + halachipedia_list,
                "customs": customs_info,
                "sources": display_sources,
                "meta": {
                    "mode": mode,
                    "community_lens": canonical_lens,
                    "source_count": 0,
                    "custom_count": len(customs_info),
                    "generated_at": int(time.time()),
                    "fallback": True,
                    "strict_blocked": True,
                    "cached": False,
                }
            }
            _set_cached_ask_payload(ask_cache_key, strict_payload)
            return jsonify(strict_payload)

        try:
            result = claude.ask_claude(
                question=question,
                sefaria_sources=flat_sources_for_claude,
                customs=customs_info,
                user_memories=user_memory_summaries,
                wiki=wiki_context_for_claude,
                halachipedia=halachipedia_list,
                mode=mode,
                community_lens=canonical_lens,
                answer_language=answer_language,
                tool_context=_build_ask_tool_context(engine),
            )

            result = _coerce_ai_answer_shape(
                result,
                question,
                mode,
                answer_language=answer_language,
            )

            result_error = str(result.get("error") or "")
            if result_error and not result_error.startswith("security_blocked"):
                raise RuntimeError(result_error or "AI request failed")

            if result_error.startswith("security_blocked"):
                blocked_answer = str(result.get("answer") or "").strip()
                if not blocked_answer:
                    blocked_answer = "Request blocked by security policy. Please submit a direct halakhic question."

                DEVTOOLS_STATS["answers_total"] += 1
                DEVTOOLS_STATS["fallback_answers"] += 1

                blocked_payload = {
                    "answer": blocked_answer,
                    "confidence": result.get("confidence", 0),
                    "wiki": [],
                    "customs": [],
                    "sources": [],
                    "meta": {
                        "mode": mode,
                        "community_lens": canonical_lens,
                        "source_count": 0,
                        "custom_count": 0,
                        "knowledge_count": len(knowledge_rows),
                        "memory_count": len(user_memory_summaries),
                        "identity_aware": bool(user_id),
                        "generated_at": int(time.time()),
                        "fallback": True,
                        "structured": False,
                        "is_prohibited": False,
                        "input_sanitized": question_was_sanitized,
                        "security": result.get("security") or {},
                        "cached": False,
                    }
                }
                _set_cached_ask_payload(ask_cache_key, blocked_payload)
                return jsonify(blocked_payload)

            structured_payload = result.get("structured")
            if not isinstance(structured_payload, dict):
                structured_payload = None

            raw_ai_answer = ""
            if structured_payload:
                raw_ai_answer = claude.render_structured_markdown(
                    structured_payload,
                    answer_language=answer_language,
                    is_simple=bool(result.get("is_simple", False)),
                )
            else:
                raw_ai_answer = str(result.get("answer") or "").strip()

            if not raw_ai_answer:
                raise RuntimeError("AI response was empty")

            needs_web_warning = use_tertiary_web_context and bool(
                wiki_context_for_claude)
            needs_internal_knowledge = (
                not has_primary_sources
                and not has_customs
                and not has_whitelisted_external
                and not wiki_context_for_claude
            )
            source_attribution_note = _build_source_attribution_note(
                has_sefaria=has_primary_sources,
                has_customs=has_customs,
                has_whitelisted_external=has_whitelisted_external,
                has_general_web=bool(wiki_context_for_claude),
                has_internal_knowledge=needs_internal_knowledge,
            )
            normalized_answer = _compose_answer_with_prefixes(
                raw_ai_answer,
                include_web_warning=needs_web_warning,
                source_attribution_note=source_attribution_note,
            )
            if not str(normalized_answer or "").strip():
                raise RuntimeError("AI response normalized to empty content")

            result["answer"] = normalized_answer
            _store_user_memory_summary(user_id, question, normalized_answer)

            DEVTOOLS_STATS["answers_total"] += 1
            display_sources = _compact_ai_sources(primary_sources)
            ai_cited = []
            if isinstance(structured_payload, dict):
                for s in (structured_payload.get("sources") or []):
                    s_str = str(s or "").strip()
                    if s_str:
                        ai_cited.append(s_str)

            # Successful AI answer path returns immediately; fallback is only for empty/error responses.
            success_payload = {
                "answer": normalized_answer,
                "confidence": result.get("confidence"),
                "wiki": wiki_list + halachipedia_list,
                "customs": customs_info,
                "sources": display_sources,
                "ai_cited_sources": ai_cited,
                "meta": {
                    "mode": mode,
                    "language": answer_language,
                    "community_lens": canonical_lens,
                    "source_count": len(primary_sources),
                    "custom_count": len(customs_info),
                    "knowledge_count": len(knowledge_rows),
                    "memory_count": len(user_memory_summaries),
                    "identity_aware": bool(user_id),
                    "generated_at": int(time.time()),
                    "fallback": bool(result.get("is_fallback", False)),
                    "structured": bool(structured_payload),
                    "is_prohibited": bool((structured_payload or {}).get("is_prohibited", False)),
                    "input_sanitized": question_was_sanitized,
                    "security": result.get("security") or {},
                    "cached": False,
                }
            }
            _set_cached_ask_payload(ask_cache_key, success_payload)
            return jsonify(success_payload)

        except Exception as ai_error:
            _capture_backend_error(
                "ask_ai_synthesis_failed",
                ai_error,
                {
                    "question": question,
                    "mode": mode,
                    "community_lens": canonical_lens,
                    "user_id": user_id or "",
                },
            )
            fallback_payload = get_halakhic_sources(question)
            fallback_warning = str(
                fallback_payload.get("warning") or "").strip()

            fallback_counts = fallback_payload.get("counts", {})
            fallback_level = str(
                fallback_payload.get("fallback_level") or "").strip().lower()
            fallback_source_note = _build_source_attribution_note(
                has_sefaria=bool(
                    fallback_counts.get("sefaria")
                    or fallback_counts.get("specific_api")
                ),
                has_customs=bool(customs_info),
                has_whitelisted_external=bool(fallback_counts.get("external")),
                has_general_web=fallback_level == "web-last-resort",
                has_internal_knowledge=fallback_level == "internal-ai-knowledge",
            )
            fallback_answer = _compose_answer_with_prefixes(
                "## Ruling\n\nAI synthesis unavailable. Returning discovered halakhic references.",
                include_web_warning=bool(fallback_warning),
                source_attribution_note=fallback_source_note,
            )

            _store_user_memory_summary(user_id, question, fallback_answer)

            DEVTOOLS_STATS["answers_total"] += 1
            DEVTOOLS_STATS["fallback_answers"] += 1
            fallback_sources = _compact_ai_sources(
                fallback_payload.get("sources", []))
            fallback_payload_response = {
                "answer": fallback_answer,
                "confidence": 0.4,
                "wiki": wiki_list + halachipedia_list,
                "customs": customs_info,
                "sources": fallback_sources,
                "meta": {
                    "mode": mode,
                    "language": answer_language,
                    "community_lens": canonical_lens,
                    "source_count": fallback_payload.get("source_count", 0),
                    "custom_count": len(customs_info),
                    "knowledge_count": len(knowledge_rows),
                    "memory_count": len(user_memory_summaries),
                    "identity_aware": bool(user_id),
                    "generated_at": int(time.time()),
                    "fallback": True,
                    "status": fallback_payload.get("status", "fallback"),
                    "fallback_detail": {
                        "keywords": fallback_payload.get("keywords", []),
                        "sequence": fallback_payload.get("sequence", []),
                        "counts": fallback_payload.get("counts", {}),
                        "level": fallback_payload.get("fallback_level", "unknown"),
                        "warning": fallback_warning,
                        "reason": str(ai_error),
                    },
                    "cached": False,
                }
            }
            # Do NOT cache AI failure/fallback responses — allow next request to retry.
            return jsonify(fallback_payload_response)

    except Exception as e:
        _capture_backend_error(
            "ask_route_critical_error",
            e,
            {
                "question": question if "question" in locals() else "",
                "mode": mode if "mode" in locals() else "",
                "community_lens": canonical_lens if "canonical_lens" in locals() else "",
            },
        )
        return jsonify({"error": "An internal error occurred while processing your request.", "detail": str(e)}), 500


@app.route("/api/stack/health")
def stack_health():
    """Return runtime readiness for Bento stack components."""
    supabase_ready = bool(_get_supabase_client())
    return jsonify({
        "flask": True,
        "vercel": True,
        "security": {
            "limiter_enabled": bool(limiter),
            "ask_limit": RATE_LIMIT_ASK,
            "global_limits": RATE_LIMIT_DEFAULT,
        },
        "clerk": {
            "configured": bool(CLERK_PUBLISHABLE_KEY and CLERK_JWT_ISSUER),
            "enforced": CLERK_ENFORCE_AUTH,
        },
        "supabase": {
            "configured": bool(SUPABASE_URL and SUPABASE_SERVICE_ROLE_KEY),
            "publishable_configured": bool(SUPABASE_URL and SUPABASE_PUBLISHABLE_KEY),
            "ready": supabase_ready,
            "prefs_table": SUPABASE_PREFS_TABLE,
        },
        "calendar": {
            "pyluach": True,
            "zmanim": True,
        },
        "external_apis": api_health.status_summary(),
        "reliability": DEVTOOLS_STATS,
    })


@app.route("/api/devtools/heartbeat")
def devtools_heartbeat():
    """Low-noise diagnostics endpoint for inspector/devtools mode."""
    started = time.time()

    checks: dict[str, Any] = {
        "clerk_configured": bool(CLERK_PUBLISHABLE_KEY and CLERK_JWT_ISSUER),
        "supabase_service_ready": bool(_get_supabase_client()),
        "supabase_publishable_ready": bool(SUPABASE_URL and SUPABASE_PUBLISHABLE_KEY),
    }

    from backend.sefaria_library import get_popular_texts
    popular_started = time.time()
    popular = get_popular_texts()
    checks["library_popular_ready"] = bool(popular)
    checks["library_popular_ms"] = int((time.time() - popular_started) * 1000)

    return jsonify({
        "ok": all(v for k, v in checks.items() if not k.endswith("_ms")),
        "ts": int(time.time()),
        "elapsed_ms": int((time.time() - started) * 1000),
        "checks": checks,
        "stats": DEVTOOLS_STATS,
    })


@app.route("/api/devtools/reliability")
def devtools_reliability():
    return jsonify({
        "stats": DEVTOOLS_STATS,
        "ts": int(time.time()),
    })


@app.route("/api/devtools/rls-audit")
def devtools_rls_audit():
    """Surface security posture for user-scoped Supabase table access."""
    has_supabase_token = bool(_extract_supabase_access_token())
    user_id = _get_request_user_id()
    return jsonify({
        "strict_rls": STRICT_SUPABASE_RLS,
        "tables": {
            "user_preferences": SUPABASE_PREFS_TABLE,
            "user_memories": SUPABASE_USER_MEMORIES_TABLE,
            "study_bookmarks": SUPABASE_STUDY_BOOKMARKS_TABLE,
        },
        "user": {
            "authenticated": bool(user_id),
            "user_id": user_id or None,
        },
        "auth": {
            "supabase_access_token_present": has_supabase_token,
            "request_scoped_client_ready": bool(_get_request_supabase_client()),
        },
        "requirement": "RLS policies should use auth.uid() = user_id for user tables.",
        "ts": int(time.time()),
    })


@app.route("/api/client-errors", methods=["POST"])
def client_errors():
    payload = request.get_json(silent=True) or {}
    context = {
        "url": str(payload.get("url") or "")[:400],
        "stack": str(payload.get("stack") or "")[:8000],
        "component": str(payload.get("component") or "")[:120],
        "user_agent": (request.headers.get("User-Agent") or "")[:320],
        "ip": _extract_client_ip() or "",
    }
    _capture_backend_error("client_error_boundary", payload.get(
        "message") or "client_error", context)
    return jsonify({"ok": True})


@app.route("/api/daily-study")
def daily_study_api():
    """Return daily refs for Daf Yomi, Rambam, and related daily study prewarming."""
    engine = get_engine()
    payload = engine.get_daily_learning() or {}
    return jsonify(payload)


@app.route("/api/devtools/segment-report", methods=["POST"])
@maybe_require_clerk_auth
def report_segment_issue():
    payload = request.get_json(silent=True) or {}
    report = {
        "ts": int(time.time()),
        "kind": (payload.get("kind") or "segment").strip()[:60],
        "message": (payload.get("message") or "").strip()[:2000],
        "segment": (payload.get("segment") or "").strip()[:160],
        "ref": (payload.get("ref") or "").strip()[:200],
        "view_type": (payload.get("view_type") or "").strip()[:40],
        "view_value": (payload.get("view_value") or "").strip()[:200],
        "client": {
            "ua": (request.headers.get("User-Agent") or "")[:300],
            "ip": _extract_client_ip() or "",
        },
    }
    claims = getattr(g, "clerk_claims", {}) or {}
    if claims.get("sub"):
        report["user_id"] = claims.get("sub")

    app.logger.warning("SEGMENT_REPORT %s",
                       json.dumps(report, ensure_ascii=True))
    DEVTOOLS_STATS["segment_reports"] += 1
    return jsonify({"ok": True, "logged": True})


@app.route("/api/auth/me")
def clerk_auth_me():
    """Returns Clerk auth status and a minimal user payload."""
    token = _extract_bearer_token()
    if not token:
        return jsonify({"authenticated": False})

    try:
        claims = _verify_clerk_token(token)
        return jsonify({
            "authenticated": True,
            "user_id": claims.get("sub"),
            "session_id": claims.get("sid"),
        })
    except Exception:
        return jsonify({"authenticated": False}), 401


@app.route("/api/user/preferences", methods=["GET", "PUT"])
@require_clerk_auth
def user_preferences():
    """Persist and fetch per-user UI preferences from Supabase."""
    claims = getattr(g, "clerk_claims", {}) or {}
    user_id = claims.get("sub")
    if not user_id:
        return jsonify({"error": "Missing user identity"}), 401

    supabase = _get_user_scoped_supabase_client()
    if not supabase and not STRICT_SUPABASE_RLS:
        supabase = _get_supabase_client()
    if not supabase:
        if STRICT_SUPABASE_RLS:
            return jsonify({"error": "Supabase authenticated session required for RLS-protected preferences."}), 403
        return jsonify({"error": "Supabase not configured"}), 503

    table = supabase.table(SUPABASE_PREFS_TABLE)

    try:
        if request.method == "GET":
            result = table.select("prefs,updated_at").eq(
                "user_id", user_id).limit(1).execute()
            rows = result.data or []
            if not rows:
                return jsonify({
                    "prefs": None,
                    "shelf": None,
                    "notes": None,
                    "reading_state": None,
                    "updated_at": None,
                })

            record = rows[0]
            if not isinstance(record, dict):
                return jsonify({
                    "prefs": None,
                    "shelf": None,
                    "notes": None,
                    "reading_state": None,
                    "updated_at": None,
                })

            stored = record.get("prefs")
            prefs = None
            shelf = None
            notes = None
            reading_state = None
            if isinstance(stored, dict):
                if any(key in stored for key in ("prefs", "shelf", "notes", "reading_state")):
                    prefs = stored.get("prefs") if isinstance(
                        stored.get("prefs"), dict) else None
                    shelf = stored.get("shelf") if isinstance(
                        stored.get("shelf"), dict) else None
                    notes = stored.get("notes") if isinstance(
                        stored.get("notes"), dict) else None
                    reading_state = stored.get("reading_state") if isinstance(
                        stored.get("reading_state"), dict) else None
                else:
                    # Legacy shape where prefs JSON was stored directly.
                    prefs = stored

            return jsonify({
                "prefs": prefs,
                "shelf": shelf,
                "notes": notes,
                "reading_state": reading_state,
                "updated_at": record.get("updated_at"),
            })

        payload = request.get_json(silent=True) or {}
        prefs = payload.get("prefs")
        if not isinstance(prefs, dict):
            return jsonify({"error": "prefs must be an object"}), 400

        shelf = payload.get("shelf")
        notes = payload.get("notes")
        reading_state = payload.get("reading_state")

        if shelf is None:
            shelf = {}
        if notes is None:
            notes = {}
        if reading_state is None:
            reading_state = {}

        if not isinstance(shelf, dict):
            return jsonify({"error": "shelf must be an object"}), 400
        if not isinstance(notes, dict):
            return jsonify({"error": "notes must be an object"}), 400
        if not isinstance(reading_state, dict):
            return jsonify({"error": "reading_state must be an object"}), 400

        now_iso = datetime.utcnow().isoformat() + "Z"
        stored_payload = {
            "prefs": prefs,
            "shelf": shelf,
            "notes": notes,
            "reading_state": reading_state,
        }
        upsert_payload = {
            "user_id": user_id,
            "prefs": stored_payload,
            "updated_at": now_iso,
        }
        table.upsert(upsert_payload).execute()
        return jsonify({"ok": True, "updated_at": now_iso})
    except Exception as e:
        _capture_backend_error("user_preferences_sync_failed", e, {
                               "user_id": str(user_id)})
        return jsonify({"error": "Failed to sync user preferences to the cloud."}), 500


@app.route("/api/bookmarks/semantic", methods=["GET", "POST"])
@require_clerk_auth
def semantic_bookmarks():
    """Persist and retrieve semantic bookmarks with notes and AI summaries."""
    claims = getattr(g, "clerk_claims", {}) or {}
    user_id = str(claims.get("sub") or "").strip()
    if not user_id:
        return jsonify({"error": "Missing user identity"}), 401

    supabase = _get_user_scoped_supabase_client()
    if not supabase and not STRICT_SUPABASE_RLS:
        supabase = _get_supabase_client()
    if not supabase:
        if STRICT_SUPABASE_RLS:
            return jsonify({"error": "Supabase authenticated session required for RLS-protected bookmarks."}), 403
        return jsonify({"error": "Supabase not configured"}), 503

    table = supabase.table(SUPABASE_STUDY_BOOKMARKS_TABLE)
    ref = ""

    try:
        if request.method == "GET":
            query_limit = _env_int("SEMANTIC_BOOKMARK_LIMIT", 50)
            result = (
                table
                .select("id,ref,label,segment_text,ai_summary,notes,created_at")
                .eq("user_id", user_id)
                .order("created_at", desc=True)
                .limit(max(1, min(query_limit, 200)))
                .execute()
            )
            return jsonify({"items": result.data or []})

        payload = request.get_json(silent=True) or {}
        ref = str(payload.get("ref") or "").strip()[:260]
        label = str(payload.get("label") or ref).strip()[:260]
        segment_text = str(payload.get("segment_text") or "").strip()[:6000]
        notes = str(payload.get("notes") or "").strip()[:3000]
        ai_summary = str(payload.get("ai_summary") or "").strip()[:3000]

        if not ref and not segment_text:
            return jsonify({"error": "A reference or text segment is required."}), 400

        summary_error = ""
        if not ai_summary and segment_text:
            summary_result = claude.summarize_with_gemini(
                segment_text, notes=notes)
            ai_summary = str(summary_result.get(
                "summary") or "").strip()[:3000]
            summary_error = str(summary_result.get("error") or "").strip()

        record = {
            "id": str(uuid4()),
            "user_id": user_id,
            "ref": ref,
            "label": label,
            "segment_text": segment_text,
            "ai_summary": ai_summary,
            "notes": notes,
            "created_at": datetime.utcnow().isoformat() + "Z",
        }

        table.insert(record).execute()
        return jsonify({
            "ok": True,
            "item": record,
            "summary_generated": bool(ai_summary),
            "summary_error": summary_error,
        })
    except Exception as e:
        _capture_backend_error("semantic_bookmark_failed", e, {
            "user_id": user_id,
            "ref": ref,
        })
        return jsonify({"error": "Failed to save semantic bookmark."}), 500


@app.route("/api/todos")
def list_todos():
    """Flask equivalent of the Next.js server query for todos."""
    supabase = _get_request_supabase_client()
    if not supabase:
        return jsonify({
            "error": "Supabase publishable client is not configured",
            "hint": "Set NEXT_PUBLIC_SUPABASE_URL and NEXT_PUBLIC_SUPABASE_PUBLISHABLE_DEFAULT_KEY",
        }), 503

    try:
        result = supabase.from_("todos").select("id,name").execute()
        return jsonify({"todos": result.data or []})
    except Exception as e:
        message = str(e)
        if "PGRST205" in message or "Could not find the table 'public.todos'" in message:
            # Treat a missing optional table as an empty list so the UI keeps working.
            return jsonify({"todos": [], "warning": "Supabase todos table is not configured"})
        return jsonify({"error": f"Failed to load todos: {message}"}), 500


@app.route("/api/library/index")
def library_index():
    """Returns report-adjusted Sefaria library tree (non-loading removals pruned, fix refs applied)."""
    from backend.sefaria_library import get_library_index
    data = get_library_index()
    return jsonify(data)


@app.route("/api/library/leaf-refs")
def library_leaf_refs():
    """Return leaf refs for a given index title to power section-grid selectors."""
    import re as _re
    from backend.sefaria_library import get_index_entry, get_index_leaf_refs

    def _parse_talmud_daf_key(daf_str):
        """Convert a daf string like '2a', '13b' to a sortable integer key."""
        m = _re.search(r'(\d+)([ab])', str(daf_str or '').strip().lower())
        if not m:
            return None
        num = int(m.group(1))
        side = 0 if m.group(2) == 'a' else 1
        return num * 2 + side

    def _extract_talmud_sections(index_title, entry):
        """
        Extract named chapter sections from a Talmud index entry's alts.Chapters.
        Returns list of {label, fromDaf, toDaf} dicts, or [] if unavailable.
        """
        alts = entry.get("alts") if isinstance(entry, dict) else None
        if not isinstance(alts, dict):
            return []
        chapters_alt = alts.get("Chapters") or alts.get("chapters")
        if not isinstance(chapters_alt, dict):
            return []
        nodes = chapters_alt.get("nodes")
        if not isinstance(nodes, list):
            return []

        sections = []
        for node in nodes:
            if not isinstance(node, dict):
                continue
            raw_title = str(node.get("title") or "").strip()
            canonical_title = entry.get("title", index_title)
            whole_ref = str(node.get("wholeRef") or "").strip()

            # Parse the daf range from wholeRef: "Berakhot 2a:1-13a:15"
            # Strip the title prefix then read the daf range
            ref_body = whole_ref
            if canonical_title and ref_body.lower().startswith(canonical_title.lower()):
                ref_body = ref_body[len(canonical_title):].lstrip(" ,")
            elif index_title and ref_body.lower().startswith(index_title.lower()):
                ref_body = ref_body[len(index_title):].lstrip(" ,")

            range_m = _re.search(
                r'(\d+[ab])(?:[\d:]*)\s*-\s*(\d+[ab])', ref_body, _re.IGNORECASE)
            if not range_m:
                continue
            from_daf = range_m.group(1).lower()
            to_daf = range_m.group(2).lower()

            # Clean the chapter label: "Chapter 1; MeEimatai" → "MeEimatai (2a–13a)"
            # Keep the human name after the semicolon, if present
            if ';' in raw_title:
                label = raw_title.split(';', 1)[1].strip()
            else:
                label = raw_title

            sections.append({
                "label": label,
                "fromDaf": from_daf,
                "toDaf": to_daf,
            })

        return sections

    def _collapse_talmud_leaf_refs(index_title, refs, max_items=260):
        """Convert segment-level Talmud refs into unique daf refs for stable grid rendering."""
        if not isinstance(refs, list) or not refs:
            return []

        normalized_title = str(index_title or "").strip()
        compact_refs = []
        seen = set()

        for ref_value in refs:
            ref_text = str(ref_value or "").strip()
            if not ref_text:
                continue

            body = ref_text
            if normalized_title and body.lower().startswith(normalized_title.lower()):
                body = body[len(normalized_title):].lstrip(" ,")

            daf_match = _re.search(r"(\d+[ab])", body, _re.IGNORECASE)
            if not daf_match:
                continue

            daf = daf_match.group(1).lower()
            if daf in seen:
                continue
            seen.add(daf)

            compact_refs.append(f"{normalized_title} {daf}".strip())
            if len(compact_refs) >= max_items:
                break

        return compact_refs

    def _synthesize_section_refs(index_title, max_items=140):
        entry = get_index_entry(index_title)
        schema = entry.get("schema", {}) if isinstance(entry, dict) else {}
        if not isinstance(schema, dict):
            return [], []

        lengths = schema.get("lengths") if isinstance(
            schema.get("lengths"), list) else []
        if not lengths:
            return [], []

        try:
            first_level_count = int(lengths[0])
        except (TypeError, ValueError):
            return [], []

        if first_level_count <= 1:
            return [], []

        section_names = schema.get("sectionNames") if isinstance(
            schema.get("sectionNames"), list) else []
        address_types = schema.get("addressTypes") if isinstance(
            schema.get("addressTypes"), list) else []

        first_section_name = str(section_names[0] or "").strip(
        ).lower() if section_names else ""
        first_address_type = str(address_types[0] or "").strip(
        ).lower() if address_types else ""

        refs = []
        if first_section_name == "daf" or first_address_type == "talmud":
            # Sefaria Talmud indexing starts at 2a.
            for idx in range(first_level_count):
                daf_num = (idx // 2) + 2
                side = "a" if idx % 2 == 0 else "b"
                refs.append(f"{index_title} {daf_num}{side}")
                if len(refs) >= max_items:
                    break
            sections = _extract_talmud_sections(index_title, entry)
            return refs, sections

        for idx in range(1, first_level_count + 1):
            refs.append(f"{index_title} {idx}")
            if len(refs) >= max_items:
                break

        return refs, []

    requested_title = _decode_route_ref(request.args.get("title", ""))
    title = str(requested_title or "").strip()
    max_refs = _coerce_int(request.args.get("max"), 140,
                           min_value=1, max_value=260)

    if not title:
        return jsonify({"title": "", "refs": [], "sections": []})

    sections = []
    try:
        refs = get_index_leaf_refs(title, max_refs=max_refs)
    except Exception:
        refs = []

    if not isinstance(refs, list):
        refs = []

    if len(refs) <= 1:
        refs, sections = _synthesize_section_refs(title, max_items=max_refs)
    else:
        # Try to extract sections even when refs came from get_index_leaf_refs
        try:
            entry = get_index_entry(title)
            sections = _extract_talmud_sections(title, entry)
        except Exception:
            sections = []

    if sections:
        collapsed_refs = _collapse_talmud_leaf_refs(
            title, refs, max_items=max_refs)
        if collapsed_refs:
            refs = collapsed_refs

    return jsonify({
        "title": title,
        "refs": refs,
        "sections": sections,
    })


@app.route("/api/library/popular")
def library_popular():
    """Returns curated popular texts per category."""
    from backend.sefaria_library import get_popular_texts
    return jsonify(get_popular_texts())


@app.route("/api/text/<path:ref>")
def get_text_inline(ref):
    """Fetches a Sefaria text inline — Hebrew + English + metadata."""
    from backend.sefaria_library import get_text
    decoded_ref = _decode_route_ref(ref)
    data = get_text(decoded_ref)

    if isinstance(data, dict) and data.get("error"):
        error_type = data.get("error_type", "")
        if error_type == "sefaria_blocked" or "blocked" in str(data.get("error", "")).lower():
            return jsonify(data), 503

    should_translate = str(request.args.get("autotranslate", "1")).strip().lower() not in {
        "0", "false", "no", "off"
    }
    if should_translate and isinstance(data, dict) and not data.get("error"):
        data = _fill_missing_english_lines(data)

    return jsonify(data)


@app.route("/api/diagnostics/sefaria")
def sefaria_diagnostics():
    """Real-time availability probe for the upstream Sefaria API.
    Returns status for both the v3 and v2 endpoints, plus cached block info.
    """
    from backend.sefaria_library import check_sefaria_availability
    result = check_sefaria_availability()
    http_status = 200 if result.get("overall_available") else 503
    return jsonify(result), http_status


@app.route("/api/word/meaning")
def get_word_meaning():
    """Look up a highlighted word meaning (best-effort for Hebrew and English)."""
    raw_word = str(request.args.get("word", "") or "").strip()
    if not raw_word:
        return jsonify({"error": "Missing word parameter"}), 400

    requested_lang = str(request.args.get(
        "lang", "en") or "en").strip().lower()
    if requested_lang not in {"en", "he"}:
        requested_lang = "en"

    word_is_hebrew = _contains_hebrew_letters(raw_word)
    if word_is_hebrew and requested_lang == "he":
        # For Hebrew source words we keep definitions in English to avoid transliteration-heavy round-trips.
        requested_lang = "en"

    if word_is_hebrew:
        meaning, source = _lookup_hebrew_word_meaning(raw_word)
    else:
        meaning, source = _lookup_english_word_meaning(raw_word)

    if meaning and requested_lang == "he" and not word_is_hebrew and not _contains_hebrew_letters(meaning):
        translated_meaning, translated_source = _translate_english_text_online(
            meaning)
        if translated_meaning:
            meaning = translated_meaning
            source_parts = [part for part in [
                source, translated_source] if part]
            source = "+".join(source_parts)

    alternatives = _collect_word_meaning_alternatives(
        raw_word=raw_word,
        primary_meaning=meaning,
        word_is_hebrew=word_is_hebrew,
    )
    if alternatives:
        meaning = alternatives[0]

    if not meaning:
        return jsonify({
            "word": raw_word,
            "meaning": "",
            "alternatives": [],
            "source": "",
            "status": "not_found",
            "lang": requested_lang,
        }), 404

    return jsonify({
        "word": raw_word,
        "meaning": meaning,
        "alternatives": alternatives,
        "source": source,
        "status": "ok",
        "lang": requested_lang,
    })


def _chapter_export_plain_text(title, ref, lines):
    header = [str(title or "").strip(), str(ref or "").strip(), ""]
    body = []
    for idx, line in enumerate(lines, start=1):
        segment = str(line.get("segment") or idx).strip()
        he = str(line.get("he") or "").strip()
        en = str(line.get("en") or "").strip()
        body.append(f"Segment {segment}")
        if he:
            body.append(f"Hebrew: {he}")
        if en:
            body.append(f"English: {en}")
        body.append("")

    return "\n".join(header + body).strip() + "\n"


def _wrap_text_for_export(text, max_chars=96):
    value = re.sub(r"\s+", " ", str(text or "").strip())
    if not value:
        return [""]

    words = value.split(" ")
    chunks = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if len(candidate) <= max_chars:
            current = candidate
            continue
        if current:
            chunks.append(current)
        current = word
    if current:
        chunks.append(current)

    return chunks or [value]


@app.route("/api/export/chapter", methods=["POST"])
@maybe_require_clerk_auth
def export_chapter():
    payload = request.get_json(silent=True) or {}
    title = str(payload.get("title") or payload.get(
        "label") or "shelah-chapter").strip()
    ref = str(payload.get("ref") or "").strip()
    export_format = str(payload.get("format") or "txt").strip().lower()
    lines = payload.get("lines") if isinstance(
        payload.get("lines"), list) else []

    if export_format not in {"txt", "docx", "pdf"}:
        return jsonify({"error": "Unsupported export format"}), 400

    normalized_lines = []
    for idx, line in enumerate(lines, start=1):
        if not isinstance(line, dict):
            continue
        he = str(line.get("he") or "").strip()
        en = str(line.get("en") or "").strip()
        if not he and not en:
            continue
        normalized_lines.append({
            "segment": str(line.get("segment") or idx).strip(),
            "he": he,
            "en": en,
        })

    if not normalized_lines:
        return jsonify({"error": "No chapter lines available to export"}), 400

    file_safe = re.sub(r"[^a-z0-9]+", "-", title.lower()
                       ).strip("-") or "shelah-chapter"
    plain_text = _chapter_export_plain_text(title, ref, normalized_lines)

    if export_format == "txt":
        txt_buffer = io.BytesIO(plain_text.encode("utf-8"))
        txt_buffer.seek(0)
        return send_file(
            txt_buffer,
            as_attachment=True,
            download_name=f"{file_safe}.txt",
            mimetype="text/plain; charset=utf-8",
        )

    if export_format == "docx":
        if Document is None:
            return jsonify({"error": "DOCX export is unavailable on this server"}), 503

        document = Document()
        document.add_heading(title or "Sh'elah Chapter", level=1)
        if ref:
            document.add_paragraph(ref)

        for idx, line in enumerate(normalized_lines, start=1):
            segment = line.get("segment") or idx
            document.add_paragraph(f"Segment {segment}")
            if line.get("he"):
                document.add_paragraph(line["he"])
            if line.get("en"):
                document.add_paragraph(line["en"])

        docx_buffer = io.BytesIO()
        document.save(docx_buffer)
        docx_buffer.seek(0)
        return send_file(
            docx_buffer,
            as_attachment=True,
            download_name=f"{file_safe}.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    if canvas is None or LETTER is None:
        return jsonify({"error": "PDF export is unavailable on this server"}), 503

    pdf_buffer = io.BytesIO()
    pdf = canvas.Canvas(pdf_buffer, pagesize=LETTER)
    width, height = LETTER
    y = height - 48
    left = 42

    def draw_line(text):
        nonlocal y
        for chunk in _wrap_text_for_export(text, max_chars=96):
            if y < 48:
                pdf.showPage()
                y = height - 48
            pdf.drawString(left, y, chunk)
            y -= 14

    draw_line(title or "Sh'elah Chapter")
    if ref:
        draw_line(ref)
    draw_line("")

    for idx, line in enumerate(normalized_lines, start=1):
        draw_line(f"Segment {line.get('segment') or idx}")
        if line.get("he"):
            draw_line(f"Hebrew: {line['he']}")
        if line.get("en"):
            draw_line(f"English: {line['en']}")
        draw_line("")

    pdf.save()
    pdf_buffer.seek(0)
    return send_file(
        pdf_buffer,
        as_attachment=True,
        download_name=f"{file_safe}.pdf",
        mimetype="application/pdf",
    )


@app.route("/api/library/search")
def library_search():
    """Full-text search across Sefaria texts with report-based removal/fix filtering."""
    from backend.sefaria_library import search_library
    query = request.args.get("q", "")
    size = _coerce_int(request.args.get("size"), 10, min_value=1, max_value=50)
    metadata_filters = _extract_search_metadata_filters()
    if not query:
        return jsonify([])
    results = search_library(
        query, size=size, metadata_filters=metadata_filters)
    return jsonify(results)


@app.route("/api/search/suggest")
def search_suggest():
    """Omnibox suggestions: texts, prayers, communities, and AI query option."""
    from backend.sefaria_library import search_library, get_liturgy_books

    query = (request.args.get("q", "") or "").strip()
    size = _coerce_int(request.args.get("size"), 8, min_value=1, max_value=20)
    metadata_filters = _extract_search_metadata_filters()
    if not query:
        return jsonify([])

    q_lower = query.lower()
    suggestions = []
    seen = set()

    def add_item(item_type, label, value, subtitle="", score=0, label_he="", subtitle_he=""):
        key = (item_type, (value or "").lower())
        if key in seen:
            return
        seen.add(key)
        suggestions.append({
            "type": item_type,
            "label": label,
            "label_he": label_he,
            "value": value,
            "subtitle": subtitle,
            "subtitle_he": subtitle_he,
            "score": score,
        })

    alias_ref = QUICK_TEXT_ALIASES.get(q_lower)
    if alias_ref:
        add_item("text", alias_ref, alias_ref, "Popular Torah alias", 100)

    for community in COMMUNITIES.keys():
        if q_lower in community.lower():
            add_item("community", community, community,
                     "Community customs", 90)

    for alias, canonical in COMMUNITY_ALIASES.items():
        if q_lower in alias and canonical in COMMUNITIES:
            add_item("community", canonical, canonical,
                     f"Community customs (matched '{alias}')", 88)

    for book in get_liturgy_books(max_items=120):
        title = book.get("title", "")
        if title and q_lower in title.lower():
            add_item("prayer", title, title, "Sefaria liturgy", 85)

    for hit in search_library(query, size=size, metadata_filters=metadata_filters):
        ref = hit.get("ref", "")
        he_ref = (hit.get("heRef", "") or "").strip()
        categories = " > ".join(hit.get("categories", [])[:3])
        if ref:
            add_item(
                "text",
                ref,
                ref,
                categories or "Sefaria text",
                70,
                label_he=he_ref or ref,
            )

    add_item("ask", f"Ask Sh'elah: {query}", query,
             "AI synthesis", 40)

    suggestions.sort(key=lambda x: x.get("score", 0), reverse=True)
    return jsonify(suggestions[:size])


@app.route("/api/text/<path:ref>/links")
def get_text_links(ref):
    """Returns all linked commentaries & parallel texts for a given ref."""
    from backend.sefaria_library import get_linked_texts
    decoded_ref = _decode_route_ref(ref)
    return jsonify(get_linked_texts(decoded_ref))


@app.route("/api/text/<path:ref>/graph")
def get_text_graph(ref):
    """Build a lightweight source graph around a text reference."""
    from backend.sefaria_library import get_linked_texts

    decoded_ref = _decode_route_ref(ref)
    links = get_linked_texts(decoded_ref)
    nodes = [{"id": decoded_ref, "label": decoded_ref, "kind": "root"}]
    edges = []
    seen = {decoded_ref}

    for category, items in (links or {}).items():
        for item in (items or [])[:14]:
            target = (item or {}).get("ref", "")
            if not target:
                continue
            if target not in seen:
                seen.add(target)
                nodes.append({
                    "id": target,
                    "label": target,
                    "kind": "linked",
                    "category": category,
                })
            edges.append({
                "source": decoded_ref,
                "target": target,
                "label": category,
            })

    return jsonify({
        "ref": ref,
        "nodes": nodes,
        "edges": edges,
    })


@app.route("/api/library/category/<path:category>")
def library_category(category):
    """Returns all books in a given Sefaria category."""
    from backend.sefaria_library import get_category_contents
    return jsonify(get_category_contents(category))


# ─── PRAYER BOOK API (Siddur Sefard - Sefardic/Mediterranean Siddur) ──────────
# Prayer content is fetched live from Sefaria refs listed in SIDDUR_SECTION_MAP.


@app.route("/api/prayers/list")
def get_prayers_list():
    """Returns all prayer books from Sefaria Liturgy plus legacy quick services."""
    from backend.sefaria_library import get_liturgy_books

    items = []
    seen = set()

    for name in SIDDUR_SECTION_MAP.keys():
        items.append({"name": name, "title": name, "source": "legacy-service"})
        seen.add(name)

    for book in get_liturgy_books(max_items=200):
        title = book.get("title")
        if title and title not in seen:
            items.append({"name": title, "title": title,
                         "source": "sefaria-liturgy"})
            seen.add(title)

    return jsonify(items)


@app.route("/api/prayer/<name>")
def get_prayer(name):
    """Returns prayer-book preview content in English and Hebrew."""
    from backend.sefaria_library import get_text

    resolved_name = (unquote(name or "") or "").strip()
    refs = _get_prayer_refs(resolved_name)
    if not refs:
        return jsonify({"error": f"Prayer '{resolved_name}' not found"}), 404

    preview = None
    for ref in refs[:12]:
        data = get_text(ref)
        if "error" not in data and (data.get("he") or data.get("en")):
            preview = data
            break

    if not preview:
        return jsonify({"error": f"Could not load prayer '{resolved_name}' from Sefaria"}), 404

    en_preview = "\n".join([l.get("en", "") for l in preview.get(
        "lines", []) if l.get("en")][:8]).strip()
    he_preview = "\n".join([l.get("he", "") for l in preview.get(
        "lines", []) if l.get("he")][:8]).strip()
    if not en_preview:
        en_preview = f"Preview available in Hebrew for {resolved_name}."
    if not he_preview:
        he_preview = f"תצוגה מקדימה זמינה באנגלית עבור {resolved_name}."

    prayer_data = {
        "en": en_preview,
        "he": he_preview,
    }

    return jsonify({
        "name": resolved_name,
        "title": resolved_name,
        "content": prayer_data,
        "languages": ["en", "he"]
    })


@app.route("/api/siddur/full/<path:prayer_name>")
def get_siddur_full(prayer_name):
    """Fetch full prayer text from Sefaria for any supported prayer service/book."""
    from backend.sefaria_library import get_text

    resolved_name = (unquote(prayer_name or "") or "").strip()
    refs = _get_prayer_refs(resolved_name)
    if not refs:
        return jsonify({"error": f"No Sefaria mapping for '{resolved_name}'"}), 404

    combined_lines = []
    for ref in refs:
        data = get_text(ref)
        if "error" not in data and (data.get("he") or data.get("en")):
            section_title = ref.split(", ")[-1] if ", " in ref else ref
            he_title = data.get("heTitle", section_title)
            combined_lines.append({
                "he": f"<strong class='text-navy'>{he_title}</strong>",
                "en": f"<strong class='text-navy'>{section_title}</strong>",
                "type": "header"
            })
            combined_lines.extend(data.get("lines", []))

    if not combined_lines:
        return jsonify({"error": "Could not fetch prayer text from Sefaria"}), 404

    return jsonify({
        "prayer": resolved_name,
        "lines": combined_lines,
        "sources": refs
    })


# ─── COMMUNITY CUSTOMS API (Merkava) ──────────────────────────────────────────

COMMUNITIES = {
    "Ashkenaz": "ashkenaz",
    "Bukharian": "bukharian",
    "Ethiopian": "ethiopian",
    "Georgian": "georgian",
    "Greek-Romaniote": "greek-romaniote",
    "Iraqi": "iraqi",
    "Kavkazi": "mountain-jewish-kavkazi",
    "Syrian": "syrian",
    "Persian": "persian",
    "Sefardic": "sefardic",
    "Turkish-Ottoman": "turkish-ottoman-sefardic",
    "Yemenite": "yemenite",
    "Moroccan": "moroccan",
    "Israeli": "sefardic",
}

COMMUNITY_ALIASES = {
    "ashkenazi": "Ashkenaz",
    "ashkenaz": "Ashkenaz",
    "sefardi": "Sefardic",
    "sephardi": "Sefardic",
    "sefardic": "Sefardic",
    "sephardic": "Sefardic",
    "iraqi": "Iraqi",
    "mizrahi": "Iraqi",
    "syrian": "Syrian",
    "yemenite": "Yemenite",
    "yemeni": "Yemenite",
    "moroccan": "Moroccan",
    "morrocan": "Moroccan",
    "israeli": "Israeli",
    "israel": "Israeli",
    "kavkazi": "Kavkazi",
    "mountain jewish": "Kavkazi",
    "mountain-jewish": "Kavkazi",
    "kavkazi jews": "Kavkazi",
    "mountain-jewish-kavkazi": "Kavkazi",
    "bukharan": "Bukharian",
    "bukharian": "Bukharian",
    "ethiopian": "Ethiopian",
    "beta israel": "Ethiopian",
    "georgian": "Georgian",
    "persian": "Persian",
    "iranian": "Persian",
    "greek": "Greek-Romaniote",
    "romaniote": "Greek-Romaniote",
    "greek-romaniote": "Greek-Romaniote",
    "turkish": "Turkish-Ottoman",
    "ottoman": "Turkish-Ottoman",
    "ottoman sefardic": "Turkish-Ottoman",
    "turkish ottoman": "Turkish-Ottoman",
    "turkish ottoman sefardic": "Turkish-Ottoman",
    "turkish-ottoman community": "Turkish-Ottoman",
    "turkish ottoman community": "Turkish-Ottoman",
    "turkish-ottoman": "Turkish-Ottoman",
    "turkish-ottoman-sefardic": "Turkish-Ottoman",
}


@app.route("/api/communities/list")
def get_communities_list():
    """Returns list of available communities."""
    communities = sorted(COMMUNITIES.keys())
    return jsonify([{"name": c} for c in communities])


@app.route("/api/community/<name>")
def get_community(name):
    """Returns community customs data."""
    resolved_name = (unquote(name or "") or "").strip()
    canonical_name = _canonicalize_community_name(resolved_name)
    if canonical_name is None:
        return jsonify({"error": f"Community '{resolved_name}' not found"}), 404

    filename = COMMUNITIES[canonical_name]
    filepath = os.path.join(os.path.dirname(__file__),
                            "customs", f"{filename}.json")

    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            data = json.load(f)

        # Extract key information for display
        identity = data.get("identity", {})
        trusted_sources = _build_trusted_custom_sources(data)

        # Extract customs from halacha_index
        customs_content = {}
        for item in data.get("halacha_index", []) if isinstance(data, dict) else []:
            if not isinstance(item, dict):
                continue
            topic = item.get("topic", "").lower()
            category = item.get("category", "").lower()
            key = f"{category}_{topic}".strip("_")
            customs_content[key] = {
                "category": category,
                "topic": topic,
                "ruling": item.get("summary", ""),
                "common_practices": item.get("common_practices", []),
                "source": item.get("source", "") or ", ".join(trusted_sources[:4])
            }

        fallback_customs = data if isinstance(data, dict) else {}

        return jsonify({
            "name": canonical_name,
            "requested_name": resolved_name,
            "heritage_id": data.get("heritage_id") if isinstance(data, dict) else None,
            "primary_origin": identity.get("primary_origin", "") if isinstance(identity, dict) else "",
            "customs": customs_content if customs_content else fallback_customs,
            "raw_data": data  # Full data available if needed
        })
    except Exception as e:
        return jsonify({"error": f"Could not load community data: {str(e)}"}), 500


@app.route("/api/community/<name>/timeline")
def get_community_timeline(name):
    """Returns a normalized community timeline for timeline view components."""
    resolved_name = (unquote(name or "") or "").strip()
    canonical_name = _canonicalize_community_name(resolved_name)
    if canonical_name is None:
        return jsonify({"error": f"Community '{resolved_name}' not found"}), 404

    filename = COMMUNITIES[canonical_name]
    filepath = os.path.join(os.path.dirname(__file__),
                            "customs", f"{filename}.json")
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            data = json.load(f)
    except Exception as e:
        return jsonify({"error": f"Could not load community data: {str(e)}"}), 500

    timeline = []

    identity = data.get("identity", {}) if isinstance(data, dict) else {}
    origin = identity.get("primary_origin") if isinstance(
        identity, dict) else ""
    if origin:
        timeline.append({
            "title": "Primary Origin",
            "description": origin,
            "approx_period": "Historic",
        })

    for key in ("timeline", "history", "historical_timeline", "migration_story"):
        value = data.get(key) if isinstance(data, dict) else None
        if isinstance(value, list):
            for item in value:
                if isinstance(item, dict):
                    timeline.append({
                        "title": str(item.get("title") or item.get("period") or key).strip()[:120],
                        "description": str(item.get("description") or item.get("event") or "").strip()[:400],
                        "approx_period": str(item.get("year") or item.get("period") or "").strip()[:80],
                    })
                elif isinstance(item, str):
                    timeline.append({
                        "title": key.replace("_", " ").title(),
                        "description": item.strip()[:400],
                        "approx_period": "",
                    })
        elif isinstance(value, str) and value.strip():
            timeline.append({
                "title": key.replace("_", " ").title(),
                "description": value.strip()[:400],
                "approx_period": "",
            })

    if not timeline:
        timeline.append({
            "title": "Tradition",
            "description": f"{canonical_name} customs are preserved through local minhagim and halachic practice.",
            "approx_period": "Ongoing",
        })

    return jsonify({
        "name": canonical_name,
        "events": timeline[:30],
    })


# ─── TEXTS INDEX (for top menu) ───────────────────────────────────────────────
@app.route("/api/texts-index")
def get_texts_index():
    """Returns complete index of browsable texts: prayers, communities, Sefaria."""
    from backend.sefaria_library import get_liturgy_books

    return jsonify({
        "siddur": {
            "title": "Sefaria Prayer Books",
            "items": [b.get("title") for b in get_liturgy_books(max_items=200)]
        },
        "merkava": {
            "title": "Community Customs (Merkava)",
            "items": list(COMMUNITIES.keys())
        },
        "sefaria": {
            "title": "Jewish Text Library",
            "items": ["Tanakh", "Mishnah", "Talmud", "Halakhah", "Kabbalah"]
        }
    })


@app.route("/api/holidays")
def get_holidays():
    """Returns Jewish holiday events for FullCalendar via Hebcal API."""
    year = request.args.get('year', str(greg_date.today().year))
    url = (
        f"https://www.hebcal.com/hebcal?v=1&cfg=json&maj=on&min=on&mod=on"
        f"&nx=on&year={year}&month=x&ss=on&s=on&mf=on&c=off&geo=none"
    )
    try:
        r = requests.get(url, timeout=5)
        data = r.json()
        if isinstance(data, dict) and data.get("error"):
            raise ValueError(data.get("error"))
        items = data.get("items", []) if isinstance(data, dict) else data
        if not isinstance(items, list):
            items = []

        events = []
        for item in items:
            if not isinstance(item, dict):
                continue

            category = str(item.get("category") or "").strip().lower()
            title_raw = item.get("title") or ""
            title_clean = _strip_leading_symbol_prefix(title_raw)
            start = item.get("date") or item.get("start")

            if not title_clean or not start:
                continue

            emoji = _holiday_emoji_for_event(title_clean, category)
            events.append({
                "title": f"{emoji} {title_clean}",
                "start": start,
                "allDay": "T" not in str(start),
                "display": "block",
                "color": _holiday_color_for_category(category),
                "textColor": "#ffffff",
            })

        return jsonify(events)
    except Exception as e:
        app.logger.warning(f"Hebcal API failed for year {year}: {str(e)}")
        fallback = _build_pyluach_holiday_events(year) or []
        if fallback:
            return jsonify(fallback)

        # Last-resort fallback to monthly zmanim events so calendar is never empty.
        try:
            engine = get_engine()
            return jsonify(engine.get_monthly_zmanim())
        except Exception:
            return jsonify({"error": "Calendar data currently unavailable", "events": []}), 503


@app.route("/api/parasha")
def get_parasha():
    """Return current weekly Parasha information for the Torah section."""
    try:
        r = requests.get("https://www.sefaria.org/api/calendars", timeout=6)
        data = r.json()

        for item in data.get("calendar_items", []):
            title_en = (item.get("title", {}) or {}).get("en", "")
            if "Parashat" in title_en or "Parasha" in title_en:
                display_en = (item.get("displayValue", {}) or {}).get("en", "")
                display_he = (item.get("displayValue", {}) or {}).get("he", "")
                ref = item.get("ref") or ""
                return jsonify({
                    "title": display_en or title_en,
                    "heTitle": display_he,
                    "ref": ref,
                    "source": "sefaria-calendars",
                })
    except Exception:
        pass

    try:
        from backend.calendar_service import calendar_engine
        parasha_name = calendar_engine.get_parasha()
        return jsonify({
            "title": parasha_name or "Parashat HaShavua",
            "heTitle": "",
            "ref": "Genesis 1",
            "source": "calendar-fallback",
        })
    except Exception:
        return jsonify({
            "title": "Parashat HaShavua",
            "heTitle": "",
            "ref": "Genesis 1",
            "source": "default-fallback",
        })


# ─── Backward-compat route aliases ────────────────────────────────────────────
# Clients (and the audit) used shorter paths before the canonical routes above
# were established. These thin shims remove the 404s without touching any logic.

@app.route("/api/health")
def api_health_alias():
    """/api/health → /api/stack/health (backward compat)."""
    return stack_health()


@app.route("/api/preferences", methods=["GET", "PUT"])
def api_preferences_alias():
    """/api/preferences → /api/user/preferences (backward compat)."""
    return user_preferences()


@app.route("/api/communities")
def api_communities_alias():
    """/api/communities → /api/communities/list (backward compat)."""
    return get_communities_list()


if __name__ == "__main__":
    port = int(os.environ.get('PORT', 5001))
    debug_mode = os.environ.get(
        'FLASK_DEBUG', '').strip().lower() in ('1', 'true')
    app.run(debug=debug_mode, host='0.0.0.0', port=port)
